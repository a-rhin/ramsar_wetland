#!/usr/bin/env python3
"""Regenerate 4 map PNGs with exact layout specs and update docx figures."""
import warnings; warnings.filterwarnings('ignore')

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import ListedColormap, BoundaryNorm
from matplotlib.ticker import MaxNLocator
from matplotlib_scalebar.scalebar import ScaleBar
import rasterio
from rasterio.mask import mask as rio_mask
from rasterio.transform import array_bounds
from rasterio.io import MemoryFile
from rasterio.features import geometry_mask
import geopandas as gpd
from shapely.geometry import mapping
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE = Path("D:/ramsar_wetlands")
OUT  = BASE / "outputs"
SITES = ['keta', 'muni']
YEARS = [1991, 2001, 2015, 2025]
COLORS = {'Vegetation': '#2d7d2d', 'Water body': '#4a90d9', 'Built up/Bareland': '#e07b39'}
NORM   = BoundaryNorm([0.5, 1.5, 2.5, 3.5], 3)
SITE_LABELS = {'keta': 'Keta Lagoon Complex', 'muni': 'Muni-Pomadze'}


def make_cmap():
    c = ListedColormap(['#2d7d2d', '#4a90d9', '#e07b39'])
    c.set_bad(color='white', alpha=0)
    return c


def load_boundary(site):
    p = BASE / f"data/{site}/area/{site}_boundary.shp"
    if not p.exists():
        p = BASE / f"data/{site}/{site}_boundary.shp"
    return gpd.read_file(p) if p.exists() else None


def load_raster(site, year, bnd_gdf=None):
    p = BASE / f"data/{site}/rasters/{site}_classified_{year}.tif"
    with rasterio.open(p) as src:
        arr  = src.read(1).astype(np.float32)
        meta = src.meta.copy()
        tf   = src.transform
        crs  = src.crs
    meta['crs'] = crs
    if int(arr.max()) <= 2:          # 0-indexed muni-style
        if bnd_gdf is not None:
            bnd    = bnd_gdf.to_crs(crs)
            geoms  = [g for g in bnd.geometry if g is not None]
            in_bnd = geometry_mask(geoms, transform=tf, out_shape=arr.shape, invert=True)
            result = np.full(arr.shape, np.nan, dtype=np.float32)
            result[in_bnd] = arr[in_bnd] + 1
            return result, meta, tf
        return np.where(arr > 0, arr + 1, np.nan).astype(np.float32), meta, tf
    arr[arr == 0] = np.nan
    return arr, meta, tf


def load_pred_2035(site):
    p = OUT / f"{site}_predicted_2035.tif"
    with rasterio.open(p) as src:
        arr  = src.read(1).astype(np.float32)
        meta = src.meta.copy()
        tf   = src.transform
        crs  = src.crs
    meta['crs'] = crs
    arr[arr == 0] = np.nan
    return arr, meta, tf


def clip_to_boundary(arr, meta, bnd_gdf):
    crs  = meta.get('crs')
    bnd  = bnd_gdf.to_crs(crs)
    geoms = [mapping(g) for g in bnd.geometry if g is not None]
    mem_meta = meta.copy()
    mem_meta.update(dtype='float32', count=1, nodata=0)
    arr_write = np.where(np.isnan(arr), 0, arr).astype(np.float32)
    with MemoryFile() as memf:
        with memf.open(**mem_meta) as ds:
            ds.write(arr_write, 1)
        with memf.open() as ds:
            out_arr, out_tf = rio_mask(ds, geoms, crop=True, nodata=0)
    data = out_arr[0].astype(float)
    data[data == 0] = np.nan
    return data, out_tf


def plot_panel(ax, data, out_tf, title, cmap, norm,
               show_xlabel=True, show_ylabel=True):
    """Plot a clipped LULC raster with coordinate grid, north arrow, scale bar."""
    masked = np.ma.masked_invalid(data)
    h, w   = data.shape
    left, bottom, right, top = array_bounds(h, w, out_tf)
    ax.set_facecolor('white')
    ax.imshow(masked, cmap=cmap, norm=norm, interpolation='nearest',
              extent=[left, right, bottom, top], origin='upper')
    ax.set_xlim(left, right)
    ax.set_ylim(bottom, top)
    if title:
        ax.set_title(title, fontsize=13, fontweight='bold', pad=5)
    ax.axis('on')
    ax.grid(True, linewidth=0.4, color='gray', alpha=0.4, linestyle='--')
    ax.tick_params(axis='both', labelsize=9)
    ax.tick_params(axis='x', rotation=45)
    ax.ticklabel_format(style='plain', axis='both')
    ax.xaxis.set_major_locator(MaxNLocator(4))
    ax.yaxis.set_major_locator(MaxNLocator(5))
    ax.set_xlabel('Easting (m)' if show_xlabel else '', fontsize=9)
    ax.set_ylabel('Northing (m)' if show_ylabel else '', fontsize=9)
    # North arrow — top-right corner of axes, outside map data
    ax.annotate('N', xy=(0.97, 0.97), xytext=(0.97, 0.87),
                xycoords='axes fraction', fontsize=10, ha='right',
                arrowprops=dict(arrowstyle='->', color='black', lw=1.5),
                fontweight='bold')
    # Scale bar
    sb = ScaleBar(1, units='m', location='lower left',
                  font_properties={'size': 8}, frameon=True,
                  color='black', box_alpha=0.7)
    ax.add_artist(sb)


# ── Generate maps ──────────────────────────────────────────────────────────────
print("=" * 60)
print("GENERATING MAP FIGURES")
print("=" * 60)

for site in SITES:
    site_label = SITE_LABELS[site]
    print(f"\n{site.upper()}:")

    bnd  = load_boundary(site)
    cmap = make_cmap()

    # Load classified rasters
    rasters, metas = {}, {}
    for year in YEARS:
        arr, meta, tf   = load_raster(site, year, bnd)
        rasters[year]   = arr
        metas[year]     = (meta, tf)
        print(f"  {year}: shape={arr.shape}")

    # Load 2035 prediction
    pred_arr, pred_meta, pred_tf = load_pred_2035(site)
    print(f"  2035 pred: shape={pred_arr.shape}")

    # ── 2×2 multipanel (1991|2001 / 2015|2025) ──────────────────────────────
    fig, axes = plt.subplots(2, 2, figsize=(12, 11))
    fig.patch.set_facecolor('white')
    plt.subplots_adjust(left=0.08, right=0.97, top=0.93, bottom=0.08,
                        wspace=0.18, hspace=0.25)

    for idx, (ax, year) in enumerate(zip(axes.flatten(), [1991, 2001, 2015, 2025])):
        row, col = divmod(idx, 2)
        arr, (meta, tf) = rasters[year], metas[year]
        clipped, clip_tf = (clip_to_boundary(arr, meta, bnd) if bnd is not None
                            else (arr, tf))
        plot_panel(ax, clipped, clip_tf, str(year), cmap, NORM,
                   show_xlabel=(row == 1), show_ylabel=(col == 0))

    patches = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
    fig.legend(handles=patches,
               labels=['Vegetation', 'Water body', 'Built up/Bareland'],
               loc='lower center', bbox_to_anchor=(0.5, 0.01),
               ncol=3, fontsize=11, markerscale=2, frameon=True, edgecolor='gray')
    fig.suptitle(f'{site_label} — LULC Maps (1991–2025)',
                 fontsize=14, fontweight='bold', y=0.97)

    png_fp = OUT / f"{site}_multipanel.png"
    pdf_fp = OUT / f"{site}_multipanel.pdf"
    fig.savefig(png_fp, dpi=300, bbox_inches='tight', facecolor='white')
    fig.savefig(pdf_fp, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  Saved: {png_fp.name} + .pdf")

    # ── 2035 standalone ──────────────────────────────────────────────────────
    pred_clipped, pred_clip_tf = (clip_to_boundary(pred_arr, pred_meta, bnd)
                                  if bnd is not None else (pred_arr, pred_tf))

    fig2, ax2 = plt.subplots(1, 1, figsize=(8, 8))
    fig2.patch.set_facecolor('white')
    plt.subplots_adjust(left=0.12, right=0.92, top=0.88, bottom=0.12)
    plot_panel(ax2, pred_clipped, pred_clip_tf, '', cmap, NORM,
               show_xlabel=True, show_ylabel=True)
    patches2 = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
    fig2.legend(handles=patches2,
                labels=['Vegetation', 'Water body', 'Built up/Bareland'],
                loc='lower center', bbox_to_anchor=(0.5, 0.01),
                ncol=3, fontsize=11, markerscale=2, frameon=True, edgecolor='gray')
    fig2.suptitle(f'{site_label} — Predicted LULC 2035 (CA-Markov)',
                  fontsize=12, fontweight='bold', y=0.97)

    pred_map_fp = OUT / f"{site}_predicted_2035_map.png"
    fig2.savefig(pred_map_fp, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close(fig2)
    print(f"  Saved: {pred_map_fp.name}")


# ── Update docx ────────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("UPDATING WORD REPORT")
print("=" * 60)

docx_path = OUT / "wetland_degradation_report.docx"
doc  = Document(str(docx_path))
paras = doc.paragraphs

replacements = [
    ("Multi-temporal LULC maps 1991–2035", "Keta Lagoon Complex",
     OUT / "keta_multipanel.png"),
    ("Multi-temporal LULC maps 1991–2035", "Muni-Pomadze",
     OUT / "muni_multipanel.png"),
]

replaced = 0
for key1, key2, img_path in replacements:
    for i, para in enumerate(paras):
        if key1 in para.text and key2 in para.text and i > 0:
            img_para = paras[i - 1]
            # Clear all XML children of the image paragraph
            p_elem = img_para._p
            for child in list(p_elem):
                p_elem.remove(child)
            # Insert new picture
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Cm(15))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            replaced += 1
            print(f"  Replaced figure: {key2} multipanel")
            break

doc.save(str(docx_path))
print(f"  Saved: {docx_path.name}  ({replaced} figure(s) replaced)")

print("\n" + "=" * 60)
print("DONE — all 4 map PNGs regenerated, docx updated")
print("=" * 60)
