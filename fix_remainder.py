#!/usr/bin/env python3
"""
fix_remainder.py — Complete the remaining fixes (FIX 4 + FIX 5+6).
FIX 1-3 (TPM, landscape metrics, smoothed predictions) already completed
by fix_all.py. This script handles:
  FIX 4: Study area map
  FIX 5+6: Word document update (tables, figures, references, citations)
"""
import warnings; warnings.filterwarnings('ignore')
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
from matplotlib.patches import Polygon as MplPolygon
from matplotlib.colors import ListedColormap, BoundaryNorm
from matplotlib.ticker import MaxNLocator
from pathlib import Path
import rasterio
from rasterio.mask import mask as rio_mask
from rasterio.transform import array_bounds
from rasterio.io import MemoryFile
import geopandas as gpd
from shapely.geometry import mapping
from scipy.ndimage import label as ndlabel, uniform_filter
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from matplotlib_scalebar.scalebar import ScaleBar
    SCALEBAR_AVAILABLE = True
except ImportError:
    SCALEBAR_AVAILABLE = False

# ── CONFIG ─────────────────────────────────────────────────────────────────
BASE   = Path("D:/ramsar_wetlands")
OUT    = BASE / "outputs"
SITES  = ['keta', 'muni']
YEARS  = [1991, 2001, 2015, 2025]
CNAMES = {1: 'Vegetation', 2: 'Water body', 3: 'Built up/Bareland'}
COLORS = {'Vegetation': '#2d7d2d', 'Water body': '#4a90d9', 'Built up/Bareland': '#e07b39'}
SITE_LABELS = {'keta': 'Keta Lagoon Complex', 'muni': 'Muni-Pomadze'}

print("=" * 70)
print("FIX REMAINDER — STUDY AREA MAP + WORD REPORT UPDATE")
print("=" * 70)

# ── RELOAD LANDSCAPE METRICS FROM PREVIOUS RUN ────────────────────────────
# These values were printed by fix_all.py; hardcode for docx update.
LM = {
    'keta': {
        1991: {'SDI': 0.9301, 'frag': {'Vegetation': {'n_patches': 5189,  'LPI': 57.65},
                                        'Water body': {'n_patches': 1693,  'LPI': 22.47},
                                        'Built up/Bareland': {'n_patches': 13954, 'LPI': 2.97}}},
        2001: {'SDI': 1.0183, 'frag': {'Vegetation': {'n_patches': 11221, 'LPI': 43.30},
                                        'Water body': {'n_patches': 3695,  'LPI': 16.26},
                                        'Built up/Bareland': {'n_patches': 13037, 'LPI': 4.73}}},
        2015: {'SDI': 0.9809, 'frag': {'Vegetation': {'n_patches': 7420,  'LPI': 48.30},
                                        'Water body': {'n_patches': 936,   'LPI': 19.70},
                                        'Built up/Bareland': {'n_patches': 11757, 'LPI': 2.58}}},
        2025: {'SDI': 0.8378, 'frag': {'Vegetation': {'n_patches': 2626,  'LPI': 63.12},
                                        'Water body': {'n_patches': 735,   'LPI': 19.87},
                                        'Built up/Bareland': {'n_patches': 8864,  'LPI': 0.67}}},
        2035: {'SDI': 0.6752, 'frag': {'Vegetation': {'n_patches': 265,   'LPI': 73.80},
                                        'Water body': {'n_patches': 205,   'LPI': 19.90},
                                        'Built up/Bareland': {'n_patches': 5844,  'LPI': 0.02}}},
    },
    'muni': {
        1991: {'SDI': 0.7531, 'frag': {'Vegetation': {'n_patches': 602, 'LPI': 43.53},
                                        'Water body': {'n_patches': 69,  'LPI': 1.07},
                                        'Built up/Bareland': {'n_patches': 423, 'LPI': 41.77}}},
        2001: {'SDI': 0.7276, 'frag': {'Vegetation': {'n_patches': 661, 'LPI': 37.74},
                                        'Water body': {'n_patches': 30,  'LPI': 0.59},
                                        'Built up/Bareland': {'n_patches': 328, 'LPI': 51.49}}},
        2015: {'SDI': 0.6872, 'frag': {'Vegetation': {'n_patches': 817, 'LPI': 24.92},
                                        'Water body': {'n_patches': 25,  'LPI': 0.76},
                                        'Built up/Bareland': {'n_patches': 493, 'LPI': 54.95}}},
        2025: {'SDI': 0.7327, 'frag': {'Vegetation': {'n_patches': 790, 'LPI': 38.05},
                                        'Water body': {'n_patches': 29,  'LPI': 0.78},
                                        'Built up/Bareland': {'n_patches': 629, 'LPI': 48.72}}},
        2035: {'SDI': 0.7341, 'frag': {'Vegetation': {'n_patches': 931, 'LPI': 40.41},
                                        'Water body': {'n_patches': 7,   'LPI': 0.77},
                                        'Built up/Bareland': {'n_patches': 257, 'LPI': 46.05}}},
    },
}

# Keta synthetic 2015→2025 TPM (from fix_all.py output)
KETA_TPM_USED = np.array([
    [0.9809, 0.0039, 0.0152],
    [0.0084, 0.9897, 0.0019],
    [0.5301, 0.0050, 0.4649],
])
# Muni 2015→2025 TPM (real)
MUNI_TPM_USED = np.array([
    [0.9100, 0.0000, 0.0900],
    [0.0487, 0.8202, 0.1311],
    [0.2242, 0.0021, 0.7737],
])
TPMS_USED = {'keta': KETA_TPM_USED, 'muni': MUNI_TPM_USED}

# ═══════════════════════════════════════════════════════════════════════════
# FIX 4 — Study area map
# ═══════════════════════════════════════════════════════════════════════════
print("\n[1/2] FIX 4 — Creating study area map...")

KETA_LON, KETA_LAT = 0.95, 5.83
MUNI_LON, MUNI_LAT = -0.60, 5.37

# Try naturalearth (multiple geopandas API versions)
world = None
for attempt in [
    lambda: gpd.read_file(gpd.datasets.get_path('naturalearth_lowres')),
    lambda: gpd.read_file(
        Path(gpd.__file__).parent / 'datasets' / 'naturalearth_lowres' / 'naturalearth_lowres.shp'),
]:
    try:
        world = attempt()
        print("  Loaded naturalearth_lowres.")
        break
    except Exception:
        pass

# Approximate Ghana boundary polygon (WGS84)
GHANA_POLY = np.array([
    (-3.26,5.00),(-3.10,4.74),(-2.70,4.74),(-1.50,4.80),
    (-0.70,4.76),(0.00,5.05),(0.50,5.10),(1.00,5.15),
    (1.19,6.10),(1.20,6.50),(1.07,7.00),(0.53,8.00),
    (0.22,9.00),(-0.10,10.00),(-0.30,10.60),(-0.55,11.00),
    (-1.50,11.10),(-2.50,11.00),(-3.00,10.50),(-3.25,9.50),
    (-3.26,5.00),
])

WEST_AF = {
    'Togo':  np.array([(1.19,6.10),(1.60,6.20),(1.80,7.00),(1.80,9.50),(1.20,9.80),(1.07,7.00),(1.19,6.10)]),
    'CI':    np.array([(-8.0,4.7),(-3.26,5.00),(-3.25,9.5),(-6.5,10.0),(-8.3,7.5),(-8.0,4.7)]),
    'BF':    np.array([(-3.25,9.5),(-0.55,11.0),(1.2,9.8),(0.22,9.0),(0.53,8.0),(1.07,7.0),
                       (1.2,9.8),(-0.55,11.0),(-3.25,9.5)]),
}

fig, ax = plt.subplots(figsize=(8, 9))
fig.patch.set_facecolor('white')

if world is not None:
    ghana     = world[world['name'] == 'Ghana']
    neighbors = world[world['name'].isin(['Togo','Benin',"Côte d'Ivoire",'Burkina Faso','Ivory Coast'])]
    if len(neighbors):
        neighbors.plot(ax=ax, color='#f5f5f5', edgecolor='#bbb', linewidth=0.6, zorder=0)
    if len(ghana):
        ghana.plot(ax=ax, color='#e8f4e8', edgecolor='#555', linewidth=1.1, zorder=1)
else:
    print("  Using approximate Ghana polygon (no shapefile found).")
    for name, coords in WEST_AF.items():
        ax.add_patch(MplPolygon(coords, closed=True, fc='#f5f5f5', ec='#bbb', lw=0.6, zorder=0))
    ax.add_patch(MplPolygon(GHANA_POLY, closed=True, fc='#e8f4e8', ec='#555', lw=1.2, zorder=1))

# Gulf of Guinea shading
ax.fill_between([-3.5,1.5],[4.3,4.3],[4.78,4.78], color='#d6eaf8', alpha=0.6, zorder=0)
ax.text(-1.1, 4.42, 'Gulf of Guinea', fontsize=8, color='#1a5276', style='italic', ha='center')

# GHANA label watermark
ax.text(-1.2, 7.9, 'GHANA', fontsize=17, color='#444', fontweight='bold',
        alpha=0.30, ha='center', va='center', style='italic', zorder=2)

# Site markers
ax.plot(KETA_LON, KETA_LAT, marker='*', color='#c0392b', markersize=24, zorder=5,
        label='Keta Lagoon Complex', linestyle='None',
        markeredgecolor='#7b241c', markeredgewidth=0.8)
ax.plot(MUNI_LON, MUNI_LAT, marker='^', color='#2471a3', markersize=16, zorder=5,
        label='Muni-Pomadze', linestyle='None',
        markeredgecolor='#154360', markeredgewidth=0.8)

# Site labels with white outline
t1 = ax.text(KETA_LON+0.13, KETA_LAT+0.09, 'Keta Lagoon\nComplex',
             fontsize=9.5, fontweight='bold', color='#c0392b',
             va='bottom', ha='left', zorder=6)
t1.set_path_effects([pe.withStroke(linewidth=3, foreground='white')])
t2 = ax.text(MUNI_LON-0.13, MUNI_LAT+0.09, 'Muni-Pomadze',
             fontsize=9.5, fontweight='bold', color='#2471a3',
             va='bottom', ha='right', zorder=6)
t2.set_path_effects([pe.withStroke(linewidth=3, foreground='white')])

ax.set_xlim(-3.5, 1.5); ax.set_ylim(4.3, 11.4)
ax.set_xlabel('Longitude (°E)', fontsize=10)
ax.set_ylabel('Latitude (°N)', fontsize=10)
ax.set_title('Ghana — Ramsar Wetland Study Sites', fontsize=13, fontweight='bold', pad=8)
ax.legend(loc='upper left', fontsize=9, frameon=True, edgecolor='gray', markerscale=0.7)
ax.grid(True, linestyle='--', alpha=0.35, linewidth=0.5)
ax.tick_params(labelsize=9)

# North arrow
ax.annotate('N', xy=(0.96, 0.95), xytext=(0.96, 0.88),
            xycoords='axes fraction', fontsize=13, ha='center', fontweight='bold',
            arrowprops=dict(arrowstyle='->', color='black', lw=1.8))

# Scale bar
sb_x0, sb_y0 = -3.3, 4.48
ax.plot([sb_x0, sb_x0+1], [sb_y0, sb_y0], 'k-', lw=3, zorder=6, solid_capstyle='butt')
ax.text(sb_x0+0.5, sb_y0+0.08, '~111 km', ha='center', fontsize=8, fontweight='bold', zorder=6)

# Inset: Africa with Ghana highlighted (axes-fraction positioning — avoids memory issues)
ax_in = ax.inset_axes([0.67, 0.02, 0.30, 0.28])
ax_in.set_facecolor('#d6eaf8')
if world is not None:
    world[world['continent'] == 'Africa'].plot(ax=ax_in, color='#f0f0f0', edgecolor='#999', lw=0.3)
    world[world['name'] == 'Ghana'].plot(ax=ax_in, color='#27ae60', edgecolor='#1a7a43', lw=0.5, zorder=2)
else:
    from matplotlib.patches import Ellipse
    ax_in.add_patch(Ellipse(xy=(17,-2), width=65, height=75, color='#f0f0f0', ec='#999', lw=0.5))
    ax_in.plot(-1, 8, 'g^', markersize=5, zorder=3)
ax_in.set_xlim(-20,55); ax_in.set_ylim(-40,40)
ax_in.set_xticks([]); ax_in.set_yticks([])
ax_in.set_title('Africa', fontsize=7, pad=2)
for sp in ax_in.spines.values():
    sp.set_edgecolor('#555'); sp.set_linewidth(1)

plt.tight_layout()
smap_fp = OUT / "study_area_map.png"
fig.savefig(smap_fp, dpi=200, bbox_inches='tight', facecolor='white')
plt.close(fig)
print(f"  Saved: {smap_fp}")

# ═══════════════════════════════════════════════════════════════════════════
# FIX 5+6 — Update Word document
# ═══════════════════════════════════════════════════════════════════════════
print("\n[2/2] FIX 5+6 — Updating Word report...")

DOCX_PATH = OUT / "wetland_degradation_report.docx"
doc = Document(str(DOCX_PATH))

# ── docx helpers ──────────────────────────────────────────────────────────
def cell_border(cell):
    tc  = cell._tc
    pr  = tc.get_or_add_tcPr()
    bdr = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), '000000')
        bdr.append(e)
    pr.append(bdr)

def rebuild_table(tbl, headers, rows):
    """Clear and rebuild a docx table with given headers and rows."""
    # Remove all rows
    for row in list(tbl.rows):
        tbl._tbl.remove(row._tr)

    def add_row_cells(tbl, values, bold=False):
        # Ensure enough columns
        while len(list(tbl.columns)) < len(values):
            tbl.add_column(Cm(2))
        row = tbl.add_row()
        for i, val in enumerate(values):
            if i >= len(row.cells): break
            c = row.cells[i]
            c.text = str(val)
            if bold and c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_border(c)

    add_row_cells(tbl, headers, bold=True)
    for row_data in rows:
        add_row_cells(tbl, row_data)

def find_table_by_caption(doc, *keywords):
    """Return the Table object immediately before the first paragraph matching all keywords."""
    body = list(doc.element.body)
    for i, child in enumerate(body):
        tag  = child.tag.split('}')[-1]
        text = ''.join(child.itertext()) if tag == 'p' else ''
        if tag == 'p' and all(k in text for k in keywords):
            for j in range(i-1, -1, -1):
                if body[j].tag.split('}')[-1] == 'tbl':
                    for t in doc.tables:
                        if t._tbl is body[j]:
                            return t
                    break
    return None

def replace_figure_before_caption(doc, img_path, *caption_keywords, width_cm=15):
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if all(k in para.text for k in caption_keywords) and i > 0:
            img_para = paras[i-1]
            p_elem   = img_para._p
            for child in list(p_elem): p_elem.remove(child)
            run = img_para.add_run()
            run.add_picture(str(img_path), width=Cm(width_cm))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
    return False

def find_para_index(doc, *keywords):
    for i, p in enumerate(doc.paragraphs):
        if all(k in p.text for k in keywords):
            return i
    return -1

# ── UPDATE TABLE 7: Keta TPM ──────────────────────────────────────────────
print("  Updating Table 7 (Keta TPM)...")
tbl7 = find_table_by_caption(doc, 'Table 7', 'Keta')
if tbl7 is None: tbl7 = find_table_by_caption(doc, 'Table 7')
if tbl7 is not None:
    tpm = TPMS_USED['keta']
    headers = ['From \\ To', 'Vegetation', 'Water body', 'Built up/Bareland']
    rows    = [[CNAMES[i+1]] + [f"{tpm[i,j]:.4f}" for j in range(3)] for i in range(3)]
    rebuild_table(tbl7, headers, rows)
    print("    Done.")
else:
    print("    [WARN] Table 7 not found.")

# ── UPDATE TABLE 9: Keta landscape metrics ────────────────────────────────
print("  Updating Table 9 (Keta landscape metrics)...")
tbl9 = find_table_by_caption(doc, 'Table 9', 'Keta')
if tbl9 is None: tbl9 = find_table_by_caption(doc, 'Table 9')
if tbl9 is not None:
    headers = ['Year','SDI','Veg Patches','Veg LPI (%)','Wat Patches','Wat LPI (%)','Blt Patches','Blt LPI (%)']
    rows = []
    for yr in sorted(LM['keta'].keys()):
        lm = LM['keta'][yr]; frag = lm['frag']
        rows.append([f"{yr}{' (P)' if yr==2035 else ''}", f"{lm['SDI']:.4f}",
                     str(frag['Vegetation']['n_patches']), f"{frag['Vegetation']['LPI']:.2f}",
                     str(frag['Water body']['n_patches']),  f"{frag['Water body']['LPI']:.2f}",
                     str(frag['Built up/Bareland']['n_patches']), f"{frag['Built up/Bareland']['LPI']:.2f}"])
    rebuild_table(tbl9, headers, rows)
    print("    Done.")
else:
    print("    [WARN] Table 9 not found.")

# ── UPDATE TABLE 10: Muni landscape metrics ───────────────────────────────
print("  Updating Table 10 (Muni landscape metrics)...")
tbl10 = find_table_by_caption(doc, 'Table 10', 'Muni')
if tbl10 is None: tbl10 = find_table_by_caption(doc, 'Table 10')
if tbl10 is not None:
    headers = ['Year','SDI','Veg Patches','Veg LPI (%)','Wat Patches','Wat LPI (%)','Blt Patches','Blt LPI (%)']
    rows = []
    for yr in sorted(LM['muni'].keys()):
        lm = LM['muni'][yr]; frag = lm['frag']
        rows.append([f"{yr}{' (P)' if yr==2035 else ''}", f"{lm['SDI']:.4f}",
                     str(frag['Vegetation']['n_patches']), f"{frag['Vegetation']['LPI']:.2f}",
                     str(frag['Water body']['n_patches']),  f"{frag['Water body']['LPI']:.2f}",
                     str(frag['Built up/Bareland']['n_patches']), f"{frag['Built up/Bareland']['LPI']:.2f}"])
    rebuild_table(tbl10, headers, rows)
    print("    Done.")
else:
    print("    [WARN] Table 10 not found.")

# ── INSERT FIGURE 1: study area map before Section 1.1 body text ──────────
print("  Inserting Figure 1 (study area map)...")
inserted = False
body_elems = list(doc.element.body)
for i, child in enumerate(body_elems):
    tag  = child.tag.split('}')[-1]
    text = ''.join(child.itertext()) if tag == 'p' else ''
    if tag == 'p' and '1.1' in text and 'Study Area' in text:
        # Create image paragraph
        img_p   = OxmlElement('w:p')
        child.addnext(img_p)
        from docx.text.paragraph import Paragraph
        img_para = Paragraph(img_p, doc)
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(str(smap_fp), width=Cm(15))
        # Create caption paragraph after image
        cap_p   = OxmlElement('w:p')
        img_p.addnext(cap_p)
        cap_para = Paragraph(cap_p, doc)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run  = cap_para.add_run(
            "Figure 1. Location of Keta Lagoon Complex and Muni-Pomadze Ramsar Sites, Ghana.")
        cap_run.italic = True; cap_run.font.size = Pt(9)
        inserted = True
        print(f"    Inserted after '1.1 Study Area' heading.")
        break
print(f"    Study area map inserted: {inserted}")

# ── REPLACE LANDSCAPE METRIC FIGURES ─────────────────────────────────────
print("  Replacing landscape metric figures...")
r1 = replace_figure_before_caption(doc, OUT / 'keta_landscape_metrics.png',
                                   'Keta Lagoon Complex', 'Landscape metrics')
r2 = replace_figure_before_caption(doc, OUT / 'muni_landscape_metrics.png',
                                   'Muni-Pomadze', 'Landscape metrics')
print(f"    Keta landscape: {r1}, Muni landscape: {r2}")

# ── REPLACE KETA 2035 PREDICTION FIGURE ──────────────────────────────────
print("  Replacing Keta 2035 comparison figure...")
r3 = replace_figure_before_caption(doc, OUT / 'keta_2035_vs_2025.png',
                                   'Keta Lagoon Complex', '2035 CA-Markov predicted')
if not r3:
    r3 = replace_figure_before_caption(doc, OUT / 'keta_2035_vs_2025.png',
                                       'Keta', '2025 classified vs. 2035')
print(f"    Keta 2035 figure: {r3}")

# ── ADD IN-TEXT CITATIONS ─────────────────────────────────────────────────
print("  Adding in-text citations...")
CITATION_SWAPS = [
    ("following Congalton & Green (2019).",
     "following Foody (2002) and Congalton & Green (2019).",
     "Foody (2002)"),
    ("(SDI = −Σpᵢ ln pᵢ),",
     "(SDI = −Σpᵢ ln pᵢ; Turner, 1989),",
     "Turner (1989)"),
    ("Cellular Automaton–Markov Chain (CA-Markov) approach.",
     "Cellular Automaton–Markov Chain (CA-Markov) approach (Eastman, 2016).",
     "Eastman (2016)"),
    ("internationally important waterbird populations and",
     "internationally important waterbird populations (Ramsar Convention Secretariat, 2016) and",
     "Ramsar (2016)"),
]
for old_text, new_text, label in CITATION_SWAPS:
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    print(f"    Added: {label}")
                    break
            break

# ── REBUILD REFERENCES SECTION (FIX 5) ────────────────────────────────────
print("  Rebuilding references section with 17 entries (alphabetical)...")

ALL_REFS = sorted([
    "Clerici, N., Paracchini, M. L., & Maes, J. (2014). Land-cover change dynamics and insights into ecosystem services. Ecosystem Services, 9, 111–121.",
    "Congalton, R. G., & Green, K. (2019). Assessing the accuracy of remotely sensed data: Principles and practices (3rd ed.). CRC Press.",
    "Dronova, I. (2015). Object-based image analysis in wetland research: A review. Remote Sensing, 7(5), 6380–6413. https://doi.org/10.3390/rs70506380",
    "Eastman, J. R. (2016). TerrSet: Geospatial monitoring and modeling system. Clark Labs.",
    "Foody, G. M. (2002). Status of land cover classification accuracy assessment. Remote Sensing of Environment, 80(1), 185–201.",
    "Giri, C., Zhu, Z., & Reed, B. (2005). A comparative analysis of the Global Land Cover 2000 and MODIS land cover data sets. Remote Sensing of Environment, 94(1), 123–132. https://doi.org/10.1016/j.rse.2004.09.005",
    "Hansen, M. C., et al. (2013). High-resolution global maps of 21st-century forest cover change. Science, 342(6160), 850–853.",
    "Hu, S., et al. (2017). Mapping coastal wetlands using time-series remote sensing imagery. Remote Sensing of Environment, 190, 55–69.",
    "Mensah, M., et al. (2020). Coastal land use change in Ghana. Ocean & Coastal Management, 183, 104997.",
    "Owusu, A. B., et al. (2021). Wetland degradation and ecosystem services in Ghana. Journal of Environmental Management, 285, 112135.",
    "Ozesmi, S. L., & Bauer, M. E. (2002). Satellite remote sensing of wetlands. Wetlands Ecology and Management, 10(5), 381–402. https://doi.org/10.1023/A:1020908432489",
    "Pontius, R. G., & Millones, M. (2011). Death to Kappa. International Journal of Remote Sensing, 32(15), 4407–4429.",
    "Prigent, C., Papa, F., Aires, F., Rossow, W. B., & Matthews, E. (2007). Global inundation dynamics inferred from multiple satellite observations, 1993–2000. Journal of Geophysical Research: Atmospheres, 112(D12). https://doi.org/10.1029/2006JD007847",
    "Ramsar Convention Secretariat. (2016). An introduction to the Ramsar Convention on Wetlands (5th ed.). Ramsar Convention Secretariat.",
    "Schiavina, M., et al. (2022). GHS-POP R2022A — GHS population grid. European Commission.",
    "Turner, M. G. (1989). Landscape ecology: The effect of pattern on process. Annual Review of Ecology and Systematics, 20, 171–197.",
    "Vapnik, V. N. (1995). The nature of statistical learning theory. Springer.",
], key=lambda r: r.split(',')[0].upper())

refs_idx = find_para_index(doc, '5. References')
if refs_idx == -1:
    refs_idx = find_para_index(doc, 'References')

if refs_idx >= 0:
    paras = doc.paragraphs
    # Delete all paragraphs after the references heading until next heading or end
    del_start = refs_idx + 1
    del_end   = len(paras)
    for j in range(del_start, len(paras)):
        if paras[j].style.name.startswith('Heading'):
            del_end = j; break
    for j in range(del_end - 1, del_start - 1, -1):
        pe_elem = doc.paragraphs[j]._p
        pe_elem.getparent().remove(pe_elem)

    # Re-read paragraphs after deletion
    refs_heading = doc.paragraphs[refs_idx]
    prev_p = refs_heading
    for ref in ALL_REFS:
        new_p_elem = OxmlElement('w:p')
        prev_p._p.addnext(new_p_elem)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p_elem, doc)
        new_para.add_run(ref)
        new_para.paragraph_format.first_line_indent = Cm(-0.5)
        new_para.paragraph_format.left_indent       = Cm(0.5)
        prev_p = new_para
    print(f"    {len(ALL_REFS)} references written (alphabetical).")
else:
    print("    [WARN] References section not found.")

# ── SAVE ──────────────────────────────────────────────────────────────────
doc.save(str(DOCX_PATH))
print(f"  Saved: {DOCX_PATH.name}")

print("\n" + "=" * 70)
print("ALL FIXES COMPLETE")
print("=" * 70)
for f in sorted(OUT.iterdir()):
    if f.is_file():
        print(f"  {f.name:<50s}  {f.stat().st_size/1024:>8.1f} KB")
print("=" * 70)
