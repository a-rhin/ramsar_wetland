import os
os.environ["PYTHONUTF8"] = "1"
import numpy as np
import rasterio
from rasterio.warp import calculate_default_transform, reproject, Resampling
from rasterio.transform import array_bounds
from rasterio.mask import mask as rio_mask
from rasterio.io import MemoryFile
from shapely.geometry import mapping
import geopandas as gpd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import BoundaryNorm, ListedColormap
from matplotlib.ticker import MaxNLocator
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

try:
    from matplotlib_scalebar.scalebar import ScaleBar
    SCALEBAR = True
except ImportError:
    SCALEBAR = False

BASE = Path("D:/ramsar_wetlands")
OUT  = BASE / "outputs"
DOCX = OUT / "wetland_degradation_report_FINAL_v3.docx"

COLORS = {"Vegetation": "#2d7d2d", "Water body": "#4a90d9", "Built up/Bareland": "#e07b39"}
CMAP   = ListedColormap(["#2d7d2d", "#4a90d9", "#e07b39"])
NORM   = BoundaryNorm([0.5, 1.5, 2.5, 3.5], 3)
YEARS  = [1991, 2001, 2015, 2025]

LM_DATA = {
    "keta": {
        "years": [1991, 2001, 2015, 2025, 2035],
        "sdi":   [0.9301, 1.0183, 0.9809, 0.8378, 0.6752],
        "Vegetation":        {"n_patches": [5189,  11221, 7420,  2626, 265],
                              "LPI":       [57.65, 43.30, 48.30, 63.12, 73.80]},
        "Water body":        {"n_patches": [1693,  3695,  936,   735,  205],
                              "LPI":       [22.47, 16.26, 19.70, 19.87, 19.90]},
        "Built up/Bareland": {"n_patches": [13954, 13037, 11757, 8864, 5844],
                              "LPI":       [2.97,  4.73,  2.58,  0.67, 0.02]},
    },
    "muni": {
        "years": [1991, 2001, 2015, 2025, 2035],
        "sdi":   [0.7531, 0.7276, 0.6872, 0.7327, 0.7341],
        "Vegetation":        {"n_patches": [602,  661,  817,  790,  931],
                              "LPI":       [43.53, 37.74, 24.92, 38.05, 40.41]},
        "Water body":        {"n_patches": [69,   30,   25,   29,   7],
                              "LPI":       [1.07,  0.59,  0.76,  0.78, 0.77]},
        "Built up/Bareland": {"n_patches": [423,  328,  493,  629,  257],
                              "LPI":       [41.77, 51.49, 54.95, 48.72, 46.05]},
    },
}

# ── FIX 2: LANDSCAPE METRICS CHARTS ─────────────────────────────────────────
print("=" * 60)
print("FIX 2: Landscape metrics charts")
for site in ["keta", "muni"]:
    lm    = LM_DATA[site]
    yrs   = lm["years"]
    ylbls = [str(y) if y != 2035 else "2035(P)" for y in yrs]
    sl    = "Keta Lagoon Complex" if site == "keta" else "Muni-Pomadze"
    fig, axes = plt.subplots(1, 3, figsize=(18, 6))
    fig.patch.set_facecolor("white")
    bc = ["#2d6a2d"] * (len(yrs) - 1) + ["#1a4f1a"]
    axes[0].bar(ylbls, lm["sdi"], color=bc, edgecolor="white")
    axes[0].set_title("Shannon Diversity Index", fontweight="bold")
    axes[0].set_xlabel("Year"); axes[0].set_ylabel("SDI")
    for i, v in enumerate(lm["sdi"]):
        axes[0].text(i, v + 0.003, f"{v:.3f}", ha="center", va="bottom", fontsize=10)
    axes[0].tick_params(axis="both", labelsize=10)
    axes[0].grid(axis="y", alpha=0.3)
    for cls, color in COLORS.items():
        axes[1].plot(ylbls, lm[cls]["n_patches"], marker="o", color=color, label=cls, lw=2)
    axes[1].set_title("Number of Patches per Class", fontweight="bold")
    axes[1].set_xlabel("Year"); axes[1].set_ylabel("Patch Count")
    axes[1].legend(fontsize=10); axes[1].grid(alpha=0.3)
    axes[1].tick_params(axis="both", labelsize=10)
    for cls, color in COLORS.items():
        axes[2].plot(ylbls, lm[cls]["LPI"], marker="s", color=color, label=cls, lw=2)
    axes[2].set_title("Largest Patch Index (%)", fontweight="bold")
    axes[2].set_xlabel("Year"); axes[2].set_ylabel("LPI (%)")
    axes[2].legend(fontsize=10); axes[2].grid(alpha=0.3)
    axes[2].tick_params(axis="both", labelsize=10)
    fig.suptitle(f"{sl} -- Landscape Metrics", fontweight="bold", fontsize=12)
    plt.tight_layout()
    fp = OUT / f"{site}_landscape_metrics.png"
    fig.savefig(fp, dpi=200, bbox_inches="tight"); plt.close(fig)
    print(f"  {fp.name}: {fp.stat().st_size//1024} KB")

# ── FIX 1: MULTIPANEL MAPS ───────────────────────────────────────────────────
print()
print("=" * 60)
print("FIX 1: Multipanel maps")
TARGET_CRS = "EPSG:32630"

def load_rst(site, year, tgt, bnd_gdf=None):
    p = BASE / f"data/{site}/rasters/{site}_classified_{year}.tif"
    if not p.exists(): return None, None, None
    with rasterio.open(p) as src:
        arr = src.read(1).astype(np.float32)
        meta = src.meta.copy(); meta["crs"] = src.crs
        sc = src.crs; stf = src.transform; sh, sw = src.height, src.width
    if int(arr.max()) <= 2:
        if bnd_gdf is not None:
            try:
                from rasterio.features import geometry_mask
                b2 = bnd_gdf.to_crs(sc)
                gs = [g for g in b2.geometry if g is not None]
                ib = geometry_mask(gs, transform=stf, out_shape=arr.shape, invert=True)
                rs = np.full(arr.shape, np.nan, dtype=np.float32)
                rs[ib] = arr[ib] + 1; arr = rs
            except Exception:
                arr = np.where(arr > 0, arr + 1, np.nan).astype(np.float32)
        else:
            arr = np.where(arr > 0, arr + 1, np.nan).astype(np.float32)
    else:
        arr[arr == 0] = np.nan
    if str(sc) != tgt:
        print(f"    Reproject {site} {year}: {sc} -> {tgt}")
        nd = np.where(np.isnan(arr), 0, arr).astype(np.float32)
        bds = array_bounds(sh, sw, stf)
        dtf, dw, dh = calculate_default_transform(sc, tgt, sw, sh, *bds)
        dst = np.zeros((dh, dw), dtype=np.float32)
        reproject(source=nd, destination=dst, src_transform=stf, src_crs=sc,
                  dst_transform=dtf, dst_crs=tgt, resampling=Resampling.nearest,
                  src_nodata=0, dst_nodata=0)
        arr = np.where(dst == 0, np.nan, dst); stf = dtf
        meta.update({"crs": rasterio.CRS.from_string(tgt),
                     "transform": dtf, "width": dw, "height": dh})
    return arr, meta, stf

def clip_bnd(arr, meta, bnd):
    crs = meta["crs"]; b2 = bnd.to_crs(crs)
    gs  = [mapping(g) for g in b2.geometry if g is not None]
    nd  = np.where(np.isnan(arr), 0, arr).astype(np.float32)
    mm  = meta.copy(); mm.update(dtype="float32", count=1, nodata=0)
    with MemoryFile() as mf:
        with mf.open(**mm) as ds: ds.write(nd, 1)
        with mf.open() as ds: out, otf = rio_mask(ds, gs, crop=True, nodata=0)
    cl = out[0].astype(float); cl[cl == 0] = np.nan
    return cl, otf

def plot_panel(ax, data, ctf, title, sx=True, sy=True):
    mk = np.ma.masked_invalid(data)
    h, w = data.shape
    l, b, r, t = array_bounds(h, w, ctf)
    ax.set_facecolor("white")
    ax.imshow(mk, cmap=CMAP, norm=NORM, interpolation="nearest",
              extent=[l, r, b, t], origin="upper")
    if title: ax.set_title(title, fontsize=13, fontweight="bold", pad=5)
    ax.axis("on")
    ax.grid(True, linewidth=0.4, color="gray", alpha=0.4, linestyle="--")
    ax.tick_params(axis="both", labelsize=9); ax.tick_params(axis="x", rotation=45)
    ax.ticklabel_format(style="plain", axis="both")
    ax.xaxis.set_major_locator(MaxNLocator(4)); ax.yaxis.set_major_locator(MaxNLocator(5))
    ax.set_xlabel("Easting (m)" if sx else "", fontsize=9)
    ax.set_ylabel("Northing (m)" if sy else "", fontsize=9)
    ax.annotate("N", xy=(0.97, 0.97), xytext=(0.97, 0.88),
                xycoords="axes fraction", fontsize=10, ha="right",
                arrowprops=dict(arrowstyle="->", color="black", lw=1.5),
                fontweight="bold", clip_on=False)
    if SCALEBAR:
        ax.add_artist(ScaleBar(1, units="m", location="lower left",
                               font_properties={"size": 8}, frameon=True,
                               color="black", box_alpha=0.7))
    return (l, b, r, t)

for site in ["keta", "muni"]:
    sl = "Keta Lagoon Complex" if site == "keta" else "Muni-Pomadze"
    print(f"  {site.upper()}...")
    bp = BASE / f"data/{site}/area/{site}_boundary.shp"
    if not bp.exists(): bp = BASE / f"data/{site}/{site}_boundary.shp"
    bnd = gpd.read_file(bp) if bp.exists() else None
    fig, axes = plt.subplots(2, 2, figsize=(12, 11))
    fig.patch.set_facecolor("white")
    plt.subplots_adjust(left=0.08, right=0.97, top=0.93, bottom=0.10,
                        wspace=0.18, hspace=0.25)
    exts = []
    for idx, (ax, yr) in enumerate(zip(axes.flatten(), YEARS)):
        row, col = divmod(idx, 2)
        res = load_rst(site, yr, TARGET_CRS, bnd)
        if res[0] is None: ax.axis("off"); continue
        arr, meta, tf = res
        cl, ctf = clip_bnd(arr, meta, bnd) if bnd is not None else (arr, tf)
        ext = plot_panel(ax, cl, ctf, str(yr), sx=(row == 1), sy=(col == 0))
        exts.append(ext)
    if exts:
        gl = min(e[0] for e in exts); gb = min(e[1] for e in exts)
        gr = max(e[2] for e in exts); gt = max(e[3] for e in exts)
        for ax in axes.flatten():
            if ax.has_data(): ax.set_xlim(gl, gr); ax.set_ylim(gb, gt)
        print(f"    Extent x=[{gl:.0f},{gr:.0f}] y=[{gb:.0f},{gt:.0f}]")
    patches = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
    fig.legend(handles=patches, labels=list(COLORS.keys()),
               loc="lower center", bbox_to_anchor=(0.5, 0.02),
               ncol=3, fontsize=11, markerscale=1.5, frameon=True, edgecolor="gray")
    fig.suptitle(f"{sl} -- LULC Maps (1991-2025)", fontsize=14, fontweight="bold", y=0.97)
    fp = OUT / f"{site}_multipanel.png"
    fig.savefig(fp, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"    Saved: {fp.name} ({fp.stat().st_size//1024} KB)")

# ── UPDATE DOCX ──────────────────────────────────────────────────────────────
print()
print("=" * 60)
print("Updating FINAL_v3 docx...")
BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
doc  = Document(str(DOCX))

def has_blip(p): return p._p.find(".//" + BLIP) is not None

def replace_img(para, img_path, w=15.0):
    pe = para._p
    for c in list(pe): pe.remove(c)
    from docx.text.paragraph import Paragraph as _P
    rp = _P(pe, doc); run = rp.add_run()
    run.add_picture(str(img_path), width=Cm(w))
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER

TARGETS = [
    ("Multi-temporal", "Keta Lagoon",     OUT / "keta_multipanel.png"),
    ("Multi-temporal", "Muni-Pomadze",    OUT / "muni_multipanel.png"),
    ("Landscape metrics", "Keta Lagoon",  OUT / "keta_landscape_metrics.png"),
    ("Landscape metrics", "Muni-Pomadze", OUT / "muni_landscape_metrics.png"),
]
for kw1, kw2, img in TARGETS:
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        t = p.text
        if kw1 in t and kw2 in t and "Figure:" in t and i > 0 and has_blip(ps[i-1]):
            replace_img(ps[i-1], img)
            print(f"  Replaced: {t[:55].encode('ascii','replace').decode()}")
            break

ps = doc.paragraphs
for i, p in enumerate(ps):
    t = p.text
    if "Table 9b: Landscape metrics" in t or "Table 10b: Landscape metrics" in t:
        if i > 0 and has_blip(ps[i-1]):
            print(f"  Removing dup image before: {t[:50].encode('ascii','replace').decode()}")
            ps[i-1]._p.getparent().remove(ps[i-1]._p)

doc.save(str(DOCX))
print(f"Saved: {DOCX.name} ({DOCX.stat().st_size/1024:.1f} KB)")