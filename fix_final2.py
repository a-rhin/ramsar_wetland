#!/usr/bin/env python3
"""
fix_final2.py — Two final fixes to wetland_degradation_report_FINAL.docx:
  FIX 1: Remove Total rows from Tables 3 & 5, add plain epoch-total sentences after footnotes
  FIX 2: Re-embed study_area_map.png as fresh inline image before Figure 1 caption
"""
import os
os.environ['PYTHONUTF8'] = '1'

from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

DOCX = Path("D:/ramsar_wetlands/outputs/wetland_degradation_report_FINAL.docx")
SMAP = Path("D:/ramsar_wetlands/outputs/study_area_map.png")

doc = Document(str(DOCX))

print("=" * 70)
print("FINAL DOCX FIX — PASS 2")
print("=" * 70)

# ════════════════════════════════════════════════════════════════════════════
# FIX 1 — Remove Total rows from Tables 3 and 5; add epoch-total sentences
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 1] Removing Total rows and adding epoch-total sentences...")

EPOCH_NOTES = {
    2: "Total classified area per epoch: approximately 1,392 km².",  # Keta
    4: "Total classified area per epoch: approximately 118 km².",    # Muni
}

body_children = list(doc.element.body)

for tbl_idx, epoch_note in EPOCH_NOTES.items():
    tbl  = doc.tables[tbl_idx]
    name = f"tables[{tbl_idx}]"

    # 1a — Remove Total row if present
    last_row = tbl.rows[-1]
    if last_row.cells[0].text.strip() == 'Total':
        tbl._tbl.remove(last_row._tr)
        print(f"  {name}: Total row removed. Rows now: {len(tbl.rows)}")
    else:
        print(f"  {name}: No Total row found (last={last_row.cells[0].text.strip()}) — skipped.")

    # Rebuild body_children list after mutation
    body_children = list(doc.element.body)

    # 1b — Find the "Note: minor rounding..." paragraph after this table
    tbl_elem = tbl._tbl
    tbl_pos  = body_children.index(tbl_elem)
    note_elem = None
    for i in range(tbl_pos + 1, min(tbl_pos + 12, len(body_children))):
        child = body_children[i]
        txt   = ''.join(child.itertext())
        if 'Note: minor rounding' in txt:
            note_elem = child
            print(f"  {name}: Found note paragraph at body pos {i}: {repr(txt[:60])}")
            break

    if note_elem is None:
        print(f"  {name}: [WARN] Note paragraph not found — epoch sentence will not be added.")
        continue

    # 1c — Check if epoch sentence already present immediately after the note
    next_elem = note_elem.getnext()
    already_there = (
        next_elem is not None
        and 'approximately' in ''.join(next_elem.itertext())
    )
    if already_there:
        print(f"  {name}: Epoch sentence already present — skipped.")
        continue

    # 1d — Insert epoch sentence paragraph immediately after the note
    new_p = OxmlElement('w:p')
    note_elem.addnext(new_p)
    from docx.text.paragraph import Paragraph as _Para
    ep_para = _Para(new_p, doc)
    run = ep_para.add_run(epoch_note)
    run.font.size = Pt(9)
    run.italic = True
    ep_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    print(f"  {name}: Epoch sentence added: {repr(epoch_note)}")

    # Rebuild again for next iteration
    body_children = list(doc.element.body)

# ════════════════════════════════════════════════════════════════════════════
# FIX 2 — Re-embed study_area_map.png before Figure 1 caption
# ════════════════════════════════════════════════════════════════════════════
print("\n[FIX 2] Embedding study_area_map.png before Figure 1 caption...")

# 2a — Validate PNG
if not SMAP.exists():
    print("  [WARN] study_area_map.png not found — regenerating...")
    SMAP = None  # trigger regeneration below

if SMAP is not None:
    with open(SMAP, 'rb') as f:
        header = f.read(8)
    if header != b'\x89PNG\r\n\x1a\n':
        print("  [WARN] study_area_map.png has invalid PNG header — regenerating...")
        SMAP = None

if SMAP is None:
    print("  Regenerating study area map...")
    import numpy as np
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyArrowPatch
    from matplotlib.offsetbox import AnchoredText

    SITES = {
        'Keta Lagoon Complex':    ( 0.95,  5.83, 'red',   '*',  200),
        'Muni-Pomadze Ramsar':    (-0.60,  5.37, 'blue',  '^',  150),
    }

    GHANA_APPROX = np.array([
        [-3.26,  5.10], [-1.06,  4.74], [ 0.57,  5.33], [ 1.20,  6.10],
        [ 1.19,  7.63], [ 0.53,  8.42], [-0.03,  9.46], [-0.22, 10.52],
        [-0.07, 11.15], [-0.56, 11.01], [-2.83, 10.60], [-3.24,  9.88],
        [-3.73,  9.37], [-3.24,  7.77], [-3.26,  5.10],
    ])

    fig, ax = plt.subplots(figsize=(8, 9))
    ax.fill(GHANA_APPROX[:, 0], GHANA_APPROX[:, 1], color='#d4e6a0', zorder=1)
    ax.plot(GHANA_APPROX[:, 0], GHANA_APPROX[:, 1], 'k-', lw=0.8, zorder=2)
    ax.set_xlim(-3.8, 1.5); ax.set_ylim(4.5, 11.5)
    ax.set_xlabel("Longitude"); ax.set_ylabel("Latitude")
    ax.set_title("Ramsar Wetland Study Sites, Ghana", fontsize=13, fontweight='bold')

    legend_handles = []
    for label, (lon, lat, col, marker, ms) in SITES.items():
        sc = ax.scatter(lon, lat, c=col, marker=marker, s=ms, zorder=5,
                        edgecolors='black', linewidths=0.5)
        ax.annotate(label, (lon, lat), textcoords='offset points',
                    xytext=(8, 4), fontsize=8, zorder=6)
        legend_handles.append(mpatches.Patch(color=col, label=label))
    ax.legend(handles=legend_handles, loc='upper left', fontsize=8)

    # North arrow
    ax.annotate('N', xy=(1.35, 11.1), fontsize=12, fontweight='bold',
                ha='center', va='center')
    ax.annotate('', xy=(1.35, 11.3), xytext=(1.35, 11.0),
                arrowprops=dict(arrowstyle='->', color='black', lw=1.5))

    # West Africa inset
    iax = ax.inset_axes([0.67, 0.02, 0.30, 0.28])
    WA_OUTLINE = np.array([
        [-18, 4], [16, 4], [16, 15], [4, 15], [-18, 15], [-18, 4]
    ])
    iax.plot(WA_OUTLINE[:,0], WA_OUTLINE[:,1], 'k-', lw=0.5)
    iax.fill(GHANA_APPROX[:,0], GHANA_APPROX[:,1], color='#8B0000', alpha=0.7)
    iax.set_xlim(-18, 16); iax.set_ylim(4, 16)
    iax.set_aspect('equal'); iax.axis('off')
    iax.set_title('West\nAfrica', fontsize=6)

    SMAP = Path("D:/ramsar_wetlands/outputs/study_area_map.png")
    plt.savefig(str(SMAP), dpi=200, bbox_inches='tight')
    plt.close()
    print(f"  Regenerated: {SMAP.stat().st_size/1024:.1f} KB")

# 2b — Find Figure 1 caption paragraph
cap_para = None
for p in doc.paragraphs:
    if 'Figure 1' in p.text and 'Keta Lagoon Complex' in p.text:
        cap_para = p
        break

if cap_para is None:
    print("  [WARN] Figure 1 caption not found — skipping image embed.")
else:
    prev_elem = cap_para._p.getprevious()

    if prev_elem is not None:
        prev_xml = etree.tostring(prev_elem, encoding='unicode')
        has_drawing = 'drawing' in prev_xml

        if has_drawing:
            # Clear existing drawing and replace with fresh image
            for child in list(prev_elem):
                prev_elem.remove(child)
            from docx.text.paragraph import Paragraph as _Para
            img_para = _Para(prev_elem, doc)
            run = img_para.add_run()
            run.add_picture(str(SMAP), width=Cm(14))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print("  Replaced existing drawing paragraph with fresh study_area_map.png image.")
        else:
            # Insert a brand new paragraph with the image before the caption
            new_p = OxmlElement('w:p')
            cap_para._p.addprevious(new_p)
            from docx.text.paragraph import Paragraph as _Para
            img_para = _Para(new_p, doc)
            run = img_para.add_run()
            run.add_picture(str(SMAP), width=Cm(14))
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print("  Inserted new image paragraph before Figure 1 caption.")
    else:
        # No previous sibling at all — insert before caption
        new_p = OxmlElement('w:p')
        cap_para._p.addprevious(new_p)
        from docx.text.paragraph import Paragraph as _Para
        img_para = _Para(new_p, doc)
        run = img_para.add_run()
        run.add_picture(str(SMAP), width=Cm(14))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("  Inserted image paragraph (no prior sibling) before Figure 1 caption.")

    # Verify
    cap_para_new = None
    for p in doc.paragraphs:
        if 'Figure 1' in p.text and 'Keta Lagoon Complex' in p.text:
            cap_para_new = p
            break
    if cap_para_new:
        prev2 = cap_para_new._p.getprevious()
        has2  = prev2 is not None and 'drawing' in etree.tostring(prev2, encoding='unicode')
        print(f"  Post-fix verification: previous sibling has drawing = {has2}")

# ════════════════════════════════════════════════════════════════════════════
# SAVE  (write to temp then rename — avoids PermissionError if Word has it open)
# ════════════════════════════════════════════════════════════════════════════
import shutil
TMP = DOCX.with_suffix('.tmp.docx')
doc.save(str(TMP))
if DOCX.exists():
    DOCX.unlink()
shutil.move(str(TMP), str(DOCX))
size_kb = DOCX.stat().st_size / 1024
print(f"\n{'='*70}")
print(f"SAVED: {DOCX.name}  ({size_kb:.1f} KB)")
print("=" * 70)
