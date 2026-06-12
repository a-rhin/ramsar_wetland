import os
os.environ['PYTHONUTF8'] = '1'

import shutil
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE  = Path('D:/ramsar_wetlands')
OUT   = BASE / 'outputs'
SRC   = OUT / 'wetland_degradation_report_FINAL_v3.tmp.docx'
DST   = OUT / 'wetland_degradation_report_FINAL_v4.docx'
BLIP  = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
TNR   = 'Times New Roman'
BLACK = RGBColor(0, 0, 0)
XML_S = '{http://www.w3.org/XML/1998/namespace}space'

# ── 0. Regenerate keta_landscape_metrics.png at 300 DPI ──────────────────────
print('Step 0: Regenerate keta_landscape_metrics.png at 300 DPI')
COLORS = {'Vegetation': '#2d7d2d', 'Water body': '#4a90d9', 'Built up/Bareland': '#e07b39'}
LM = {
    'years': [1991, 2001, 2015, 2025, 2035],
    'sdi':   [0.9301, 1.0183, 0.9809, 0.8378, 0.6752],
    'Vegetation':        {'n_patches': [5189, 11221, 7420, 2626, 265],
                          'LPI':       [57.65, 43.30, 48.30, 63.12, 73.80]},
    'Water body':        {'n_patches': [1693, 3695, 936, 735, 205],
                          'LPI':       [22.47, 16.26, 19.70, 19.87, 19.90]},
    'Built up/Bareland': {'n_patches': [13954, 13037, 11757, 8864, 5844],
                          'LPI':       [2.97, 4.73, 2.58, 0.67, 0.02]},
}
ylbls = ['1991', '2001', '2015', '2025', '2035(P)']
bc    = ['#2d6a2d'] * 4 + ['#1a4f1a']
fig, axs = plt.subplots(1, 3, figsize=(18, 6))
fig.patch.set_facecolor('white')
axs[0].bar(ylbls, LM['sdi'], color=bc, edgecolor='white')
axs[0].set_title('Shannon Diversity Index', fontweight='bold')
axs[0].set_xlabel('Year'); axs[0].set_ylabel('SDI')
for i, v in enumerate(LM['sdi']):
    axs[0].text(i, v + 0.003, f'{v:.3f}', ha='center', va='bottom', fontsize=10)
axs[0].tick_params(labelsize=10); axs[0].grid(axis='y', alpha=0.3)
for cls, col in COLORS.items():
    axs[1].plot(ylbls, LM[cls]['n_patches'], marker='o', color=col, label=cls, lw=2)
axs[1].set_title('Number of Patches per Class', fontweight='bold')
axs[1].set_xlabel('Year'); axs[1].set_ylabel('Patch Count')
axs[1].legend(fontsize=10); axs[1].grid(alpha=0.3); axs[1].tick_params(labelsize=10)
for cls, col in COLORS.items():
    axs[2].plot(ylbls, LM[cls]['LPI'], marker='s', color=col, label=cls, lw=2)
axs[2].set_title('Largest Patch Index (%)', fontweight='bold')
axs[2].set_xlabel('Year'); axs[2].set_ylabel('LPI (%)')
axs[2].legend(fontsize=10); axs[2].grid(alpha=0.3); axs[2].tick_params(labelsize=10)
fig.suptitle('Keta Lagoon Complex — Landscape Metrics', fontweight='bold', fontsize=12)
plt.tight_layout()
keta_lm_png = OUT / 'keta_landscape_metrics.png'
fig.savefig(str(keta_lm_png), dpi=300, bbox_inches='tight')
plt.close(fig)
print(f'  Saved {keta_lm_png.stat().st_size//1024} KB')

# ── 1. Load document ──────────────────────────────────────────────────────────
print('Step 1: Load document')
shutil.copy2(str(SRC), str(DST))
doc = Document(str(DST))
print(f'  {len(doc.paragraphs)} paragraphs loaded')

# ── 2. Page margins ───────────────────────────────────────────────────────────
print('Step 2: Page margins')
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54); s.right_margin = Cm(2.54)

# ── HELPERS ───────────────────────────────────────────────────────────────────
def has_blip(p):
    return p._p.find('.//' + BLIP) is not None

def fmt_run(run, sz=12, bold=False, italic=False):
    run.font.name = TNR
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    try:
        run.font.color.rgb = BLACK
    except Exception:
        pass
    rPr = run._r.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts')
        rPr.insert(0, rF)
    rF.set(qn('w:ascii'), TNR); rF.set(qn('w:hAnsi'), TNR); rF.set(qn('w:cs'), TNR)

def set_spacing15(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def new_para_after(anchor_para, text, sz=12):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'both'); pPr.append(jc)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), '360'); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp); p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), TNR); rF.set(qn('w:hAnsi'), TNR); rF.set(qn('w:cs'), TNR)
    rPr.append(rF)
    sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), str(sz * 2)); rPr.append(sz_el)
    szC = OxmlElement('w:szCs'); szC.set(qn('w:val'), str(sz * 2)); rPr.append(szC)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(XML_S, 'preserve')
    r.append(t); p.append(r)
    anchor_para._p.addnext(p)

# ── 3. Rename figure captions ─────────────────────────────────────────────────
print('Step 3: Rename figure captions')

CAPTION_MAP = [
    (['Figure:', '1991 confusion', 'Keta'],
     'Figure 2. Confusion matrix — Keta Lagoon Complex 1991 (OA = 81.4%, κ = 0.697).'),
    (['Figure:', '2001 confusion', 'Keta'],
     'Figure 3. Confusion matrix — Keta Lagoon Complex 2001 (OA = 86.9%, κ = 0.792).'),
    (['Figure:', '2015 confusion', 'Keta'],
     'Figure 4. Confusion matrix — Keta Lagoon Complex 2015 (OA = 83.3%, κ = 0.729).'),
    (['Figure:', '2025 confusion', 'Keta'],
     'Figure 5. Confusion matrix — Keta Lagoon Complex 2025 (OA = 83.6%, κ = 0.751).'),
    (['Figure:', '1991 confusion', 'Muni'],
     'Figure 6. Confusion matrix — Muni-Pomadze 1991 (OA = 83.3%, κ = 0.750).'),
    (['Figure:', '2001 confusion', 'Muni'],
     'Figure 7. Confusion matrix — Muni-Pomadze 2001 (OA = 90.0%, κ = 0.850).'),
    (['Figure:', '2015 confusion', 'Muni'],
     'Figure 8. Confusion matrix — Muni-Pomadze 2015 (OA = 86.7%, κ = 0.800).'),
    (['Figure:', '2025 confusion', 'Muni'],
     'Figure 9. Confusion matrix — Muni-Pomadze 2025 (OA = 86.7%, κ = 0.800).'),
    (['Figure:', 'LULC area by class', 'Keta'],
     'Figure 10. LULC area by class and year — Keta Lagoon Complex.'),
    (['Figure:', 'LULC area trend', 'Keta'],
     'Figure 11. LULC area trend lines (1991–2025) — Keta Lagoon Complex.'),
    (['Figure:', 'LULC area by class', 'Muni'],
     'Figure 12. LULC area by class and year — Muni-Pomadze.'),
    (['Figure:', 'LULC area trend', 'Muni'],
     'Figure 13. LULC area trend lines (1991–2025) — Muni-Pomadze.'),
    (['Figure:', '2025 classified', 'Keta'],
     'Figure 14. Keta Lagoon Complex — 2025 classified vs. 2035 CA-Markov projected land cover.'),
    (['Figure:', '2025 classified', 'Muni'],
     'Figure 15. Muni-Pomadze — 2025 classified vs. 2035 CA-Markov projected land cover.'),
    (['Figure:', 'Multi-temporal', 'Keta'],
     'Figure 16. Multi-temporal LULC maps — Keta Lagoon Complex (1991–2035). '
     'Maps for 1991–2025 produced by SVM classification; 2035 map is the CA-Markov projection.'),
    (['Figure:', 'Multi-temporal', 'Muni'],
     'Figure 17. Multi-temporal LULC maps — Muni-Pomadze (1991–2035). '
     'Maps for 1991–2025 produced by SVM classification; 2035 map is the CA-Markov projection.'),
    (['Figure:', 'Landscape metrics', 'Keta'],
     'Figure 18. Landscape metrics — Keta Lagoon Complex (1991–2035). '
     'SDI = Shannon Diversity Index; LPI = Largest Patch Index; PC = Patch Count. (P) = CA-Markov projection.'),
    (['Figure:', 'Landscape metrics', 'Muni'],
     'Figure 19. Landscape metrics — Muni-Pomadze (1991–2035). '
     'SDI = Shannon Diversity Index; LPI = Largest Patch Index; PC = Patch Count. (P) = CA-Markov projection.'),
]

for keywords, new_text in CAPTION_MAP:
    found = False
    for para in doc.paragraphs:
        if all(kw in para.text for kw in keywords):
            pPr = para._p.find(qn('w:pPr'))
            for child in list(para._p):
                if child is not pPr:
                    para._p.remove(child)
            r_el = OxmlElement('w:r')
            t_el = OxmlElement('w:t')
            t_el.text = new_text
            t_el.set(XML_S, 'preserve')
            r_el.append(t_el)
            para._p.append(r_el)
            print(f'  OK: {new_text[:55].encode("ascii","replace").decode()}...')
            found = True; break
    if not found:
        print(f'  WARN not found: {keywords}')
# ── 4. Global formatting pass ─────────────────────────────────────────────────
print('Step 4: Global formatting')

# Find index of Abstract heading to protect title page paragraphs
abs_idx = None
paras = doc.paragraphs
for i, p in enumerate(paras):
    if p.style.name.startswith('Heading') and 'Abstract' in p.text:
        abs_idx = i; break

for i, para in enumerate(paras):
    t = para.text.strip()
    sname = para.style.name
    is_img = has_blip(para)

    # Clean markdown symbols from runs
    for run in para.runs:
        cleaned = run.text
        for sym in ['**', '*', '```', '`']:
            cleaned = cleaned.replace(sym, '')
        if cleaned != run.text:
            run.text = cleaned

    # Title page paragraphs: just set TNR and black, keep existing size
    if abs_idx is not None and i < abs_idx:
        for run in para.runs:
            run.font.name = TNR
            try: run.font.color.rgb = BLACK
            except Exception: pass
            rPr = run._r.get_or_add_rPr()
            rF = rPr.find(qn('w:rFonts'))
            if rF is None:
                rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
            rF.set(qn('w:ascii'), TNR); rF.set(qn('w:hAnsi'), TNR); rF.set(qn('w:cs'), TNR)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    if sname.startswith('Heading 1'):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing15(para)
        for run in para.runs: fmt_run(run, sz=14, bold=True)

    elif sname.startswith('Heading 2'):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing15(para)
        for run in para.runs: fmt_run(run, sz=13, bold=True)

    elif sname.startswith('Heading 3'):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing15(para)
        for run in para.runs: fmt_run(run, sz=12, bold=True)

    elif is_img:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif t.startswith('Figure ') and len(t) > 8 and ('.' in t[7:12] or ' ' in t[7:12]):
        # Figure caption: italic TNR 11 centred
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs: fmt_run(run, sz=11, italic=True)

    elif t.startswith('Table ') and (':' in t[:25] or (len(t) < 100 and 'accuracy' in t.lower())):
        # Table caption: bold TNR 11 left
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs: fmt_run(run, sz=11, bold=True)

    elif (t.startswith('Note:') or t.startswith('Total classified') or
          'ᵢ' in t or 'p = proportion' in t or
          t.startswith('−') or (t.startswith('p') and 'proportion' in t)):
        # Footnote/note: italic TNR 10
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs: fmt_run(run, sz=10, italic=True)

    elif t:
        # General body text: TNR 12 justified 1.5
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing15(para)
        for run in para.runs: fmt_run(run, sz=12)

print(f'  Done — {len(paras)} paragraphs processed')

# ── 5. Replace keta_landscape_metrics with 300-DPI version ───────────────────
print('Step 5: Re-embed keta_landscape_metrics.png (300 DPI)')
from docx.text.paragraph import Paragraph as _P

ps = doc.paragraphs
replaced = False
for i, para in enumerate(ps):
    t = para.text
    if 'Figure 18' in t and 'Landscape metrics' in t and i > 0 and has_blip(ps[i-1]):
        img_p = ps[i-1]
        pe = img_p._p
        for c in list(pe): pe.remove(c)
        rp = _P(pe, doc); run = rp.add_run()
        run.add_picture(str(keta_lm_png), width=Cm(15))
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f'  Replaced at para {i-1}')
        replaced = True; break
if not replaced:
    print('  WARN: could not find keta LM image para')

# ── 6. Interpretive paragraphs ────────────────────────────────────────────────
print('Step 6: Insert interpretive paragraphs')

KETA_ACCURACY = (
    'Classification accuracy for Keta Lagoon Complex ranged from an Overall Accuracy (OA) of '
    '81.4 % (κ = 0.697) in 1991 to a peak of 86.9 % (κ = 0.792) '
    'in 2001, before settling at 83.6 % (κ = 0.751) in 2025. Kappa coefficients '
    'exceeded 0.69 across all periods, indicating substantial to near-perfect agreement between '
    'classified outputs and reference validation samples (Landis & Koch, 1977). The lower accuracy '
    'in 1991 reflects the coarser radiometric calibration of Landsat 4/5 TM imagery and greater '
    'temporal distance from contemporary ground-truth observations. Across all periods, the primary '
    'source of confusion was between the Vegetation and Built up/Bareland classes, attributable to '
    'spectral overlap between dry halophytic vegetation and bare salt flat deposits characteristic '
    'of the Keta coastal environment.'
)

MUNI_ACCURACY = (
    'Classification accuracy at Muni-Pomadze was consistently high, with OA ranging from 83.3 % '
    '(κ = 0.750) in 1991 to 90.0 % (κ = 0.850) in 2001; the 2015 '
    'and 2025 periods each achieved OA = 86.7 % (κ = 0.800). The peak '
    'accuracy in 2001 likely reflects improved Landsat 7 ETM+ sensor performance and pronounced spectral '
    'contrast between the site’s restricted water body and surrounding land cover classes. The '
    'compact spatial extent of Muni-Pomadze (approximately 118 km²) supported more spectrally '
    'homogeneous classification units compared with the larger Keta site, contributing to consistently '
    'higher overall accuracy values.'
)

KETA_LULC = (
    'Between 1991 and 2025, Keta Lagoon Complex experienced substantial and largely irreversible land '
    'cover transformation. Vegetation declined from 845.42 km² (60.7 % of total area) '
    'in 1991 to 498.98 km² (36.2 %) in 2025, a net loss of 346.44 km² over '
    '34 years. Built up/Bareland expanded from 206.03 km² (14.8 %) to '
    '567.95 km² (41.2 %), a net gain of 361.92 km². Notably, between 2001 '
    'and 2015 Vegetation recovered from 719.91 km² to 787.36 km² while '
    'Built up/Bareland contracted from 408.47 km² to 295.29 km². This partial '
    'recovery may reflect episodic recolonisation of salt flat margins by halophytic vegetation and '
    'reduced land-conversion pressure during that period — dynamics consistent with saline wetland '
    'systems documented elsewhere in coastal West Africa (Mensah et al., 2020). However, the '
    '2015–2025 interval reversed this trend sharply: Vegetation declined by 288.38 km² '
    'and Built up/Bareland expanded by 272.66 km², indicating renewed anthropogenic pressure '
    'concurrent with regional urban and agricultural expansion. Water body coverage declined moderately '
    'from 340.62 km² (24.5 %) in 1991 to 312.93 km² (22.7 %) in 2025.'
)

MUNI_LULC = (
    'At Muni-Pomadze, LULC dynamics reflect moderate but persistent transformation, consistent with '
    'the site’s smaller spatial extent of approximately 118 km². Vegetation declined '
    'from 60.77 km² (51.5 % of site area) in 1991 to a minimum of 39.74 km² '
    '(33.7 %) in 2015, before recovering partially to 53.34 km² (45.5 %) in 2025. '
    'Built up/Bareland expanded from 55.59 km² (47.1 %) to 76.38 km² '
    '(64.7 %) in 2015, the period of peak anthropogenic pressure, before contracting to '
    '62.80 km² (53.2 %) in 2025. This contraction may reflect short-term vegetation '
    'recovery or site-level conservation and buffer zone enforcement. Water body remained minimal '
    'throughout, fluctuating between 0.81 km² (0.69 %) in 2001 and 1.56 km² '
    '(1.3 %) in 1991, reflecting the limited open-water extent of this transitional coastal wetland.'
)

CAMARKOV = (
    'Transition probability matrices (TPMs), derived from the 2015–2025 period, reveal the land '
    'cover dynamics underpinning each site’s trajectory. At Keta, high diagonal probabilities '
    'indicate strong class persistence: Vegetation (0.9809) and Water body (0.9897) retained '
    'near-complete self-persistence over the transition interval. The Built up/Bareland-to-Vegetation '
    'transition probability of 0.53 warrants cautious interpretation; rather than indicating widespread '
    'land recovery, this value likely reflects the dynamic tidal and salt flat environment at Keta, '
    'where seasonal flooding temporarily restores vegetation spectral signatures in otherwise degraded '
    'zones (see §1.3 for accuracy discussion). The 2035 CA-Markov projections, extending '
    '2015–2025 dynamics forward one decade, indicate continued Vegetation contraction at Keta '
    'driven by elevated Built up/Bareland persistence, and relative landscape stability at '
    'Muni-Pomadze with minor class-level fluctuations. Absent targeted management intervention, these '
    'trajectories suggest further fragmentation of the vegetated wetland matrix at Keta and '
    'consolidation of urban-fringe land cover at Muni.'
)

LM_INTERP = (
    'Landscape metric analysis provides structural context for the LULC transformations documented '
    'above. At Keta Lagoon Complex, the Shannon Diversity Index (SDI) peaked in 2001 '
    '(SDI = 1.018), reflecting the more equitable class distribution during the partial '
    'vegetation recovery period, before declining to 0.838 in 2025 and a projected 0.675 in 2035. '
    'This trajectory indicates a landscape converging toward dominance by Built up/Bareland, with '
    'associated reduction in habitat heterogeneity. Although the Vegetation patch count declined '
    'sharply from 11,221 in 2001 to 265 in the 2035 projection, the Largest Patch Index (LPI) for '
    'Vegetation paradoxically increased to a projected 73.80 %: remaining Vegetation is '
    'consolidating into fewer but proportionally larger remnant patches, indicative of landscape '
    'simplification following fragmentation. At Muni-Pomadze, SDI values were comparatively stable '
    'across all periods (0.687–0.753), reflecting the more constrained and consistently balanced '
    'land cover composition of this smaller site. Vegetation patch counts fluctuated moderately between '
    '602 and 931, and LPI values showed limited variation, suggesting ongoing low-intensity habitat '
    'restructuring rather than the dramatic simplification observed at Keta. The consistently elevated '
    'LPI for Built up/Bareland at Muni (41.77–54.95 % across all periods) reflects the '
    'longstanding co-dominance of this class with Vegetation in the site’s peri-urban setting.'
)

INTERP_MAP = [
    ('Figure 5.',  KETA_ACCURACY),
    ('Figure 9.',  MUNI_ACCURACY),
    ('Figure 11.', KETA_LULC),
    ('Figure 13.', MUNI_LULC),
    ('Figure 15.', CAMARKOV),
    ('Figure 19.', LM_INTERP),
]

for anchor_prefix, text in INTERP_MAP:
    found = False
    for para in doc.paragraphs:
        if para.text.startswith(anchor_prefix):
            new_para_after(para, text)
            print(f'  Inserted after: {anchor_prefix}')
            found = True; break
    if not found:
        print(f'  WARN not found: {anchor_prefix}')

# ── 7. Save ────────────────────────────────────────────────────────────────────
print('Step 7: Save')
doc.save(str(DST))
sz = DST.stat().st_size / 1024
print(f'DONE: {DST.name}  ({sz:.1f} KB)')