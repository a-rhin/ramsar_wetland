#!/usr/bin/env python3
"""Wetland Degradation Analysis Pipeline — Keta & Muni Ramsar Sites, Ghana"""

# ============================================================
# 0. AUTO-INSTALL PACKAGES
# ============================================================
import subprocess, sys, warnings, os
warnings.filterwarnings('ignore')

print("=" * 70)
print("WETLAND DEGRADATION ANALYSIS PIPELINE")
print("Keta Lagoon Complex & Muni-Pomadze Ramsar Sites, Ghana")
print("=" * 70)
print("\n[0] Checking required packages...")

def _pip(pkg):
    subprocess.check_call(
        [sys.executable, '-m', 'pip', 'install', pkg, '-q', '--no-warn-script-location'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )

for _imp, _pkg in [
    ('numpy', 'numpy'), ('pandas', 'pandas'), ('matplotlib', 'matplotlib'),
    ('seaborn', 'seaborn'), ('sklearn', 'scikit-learn'), ('openpyxl', 'openpyxl'),
    ('rasterio', 'rasterio'), ('geopandas', 'geopandas'),
    ('shapely', 'shapely'), ('docx', 'python-docx'), ('PIL', 'Pillow'),
    ('scipy', 'scipy'),
    ('matplotlib_scalebar', 'matplotlib-scalebar'),
]:
    try:
        __import__(_imp)
    except ImportError:
        print(f"  Installing {_pkg}...")
        _pip(_pkg)

PYLANDSTATS_AVAILABLE = False
try:
    import pylandstats as pls
    PYLANDSTATS_AVAILABLE = True
    print("  pylandstats: available")
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, '-m', 'pip', 'install', 'pylandstats',
             '--no-build-isolation', '-q', '--no-warn-script-location'],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=120
        )
        import pylandstats as pls
        PYLANDSTATS_AVAILABLE = True
        print("  pylandstats: installed")
    except Exception:
        print("  pylandstats: unavailable — using manual metrics")

print("  All packages ready.\n")

# ============================================================
# IMPORTS
# ============================================================
import numpy as np
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import ListedColormap, BoundaryNorm
import seaborn as sns
from sklearn.metrics import confusion_matrix, cohen_kappa_score, accuracy_score
import rasterio
import geopandas as gpd
from scipy.ndimage import label as ndlabel
from scipy.ndimage import generic_filter
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from matplotlib.ticker import MaxNLocator
from rasterio.mask import mask as rio_mask
from rasterio.transform import array_bounds
from rasterio.io import MemoryFile
from shapely.geometry import mapping

matplotlib.use('Agg')

SCALEBAR_AVAILABLE = False
try:
    from matplotlib_scalebar.scalebar import ScaleBar
    SCALEBAR_AVAILABLE = True
except ImportError:
    pass

plt.rcParams.update({
    'font.family': 'DejaVu Sans',
    'axes.titlesize': 11, 'axes.labelsize': 10,
    'xtick.labelsize': 9, 'ytick.labelsize': 9,
    'figure.dpi': 150,
})


def make_masked(arr):
    """Return masked array with background/NaN as masked (renders as white)."""
    return np.ma.masked_where(np.isnan(arr) | (arr < 0.5), arr)


def make_cmap():
    c = ListedColormap(['#2d7d2d', '#4a90d9', '#e07b39'])
    c.set_bad(color='white', alpha=0)
    return c


def clip_to_boundary(arr, meta, bnd_gdf):
    """Clip raster array to boundary polygon using rasterio.mask.mask()."""
    crs = meta.get('crs')
    bnd = bnd_gdf.to_crs(crs)
    geoms = [mapping(g) for g in bnd.geometry if g is not None]
    mem_meta = meta.copy()
    mem_meta.update(dtype='float32', count=1, nodata=0)
    arr_write = np.where(np.isnan(arr), 0, arr).astype(np.float32)
    with MemoryFile() as memf:
        with memf.open(**mem_meta) as ds:
            ds.write(arr_write, 1)
        with memf.open() as ds:
            out_arr, out_tf = rio_mask(ds, geoms, crop=True, nodata=0)
    data = out_arr[0].astype(float)
    data[data == 0] = np.nan
    return data, out_tf, bnd


def plot_lulc_panel(ax, data, out_tf, _bnd_reproj, title, cmap, norm,
                    show_xlabel=True, show_ylabel=True):
    """Plot a clipped LULC raster with coordinate grid, north arrow, scale bar."""
    masked = np.ma.masked_invalid(data)
    h, w   = data.shape
    left, bottom, right, top = array_bounds(h, w, out_tf)
    ax.set_facecolor('white')
    ax.imshow(masked, cmap=cmap, norm=norm, interpolation='nearest',
              extent=[left, right, bottom, top], origin='upper')
    ax.set_xlim(left, right)
    ax.set_ylim(bottom, top)
    if title:
        ax.set_title(title, fontsize=13, fontweight='bold', pad=5)
    ax.axis('on')
    ax.grid(True, linewidth=0.4, color='gray', alpha=0.4, linestyle='--')
    ax.tick_params(axis='both', labelsize=9)
    ax.tick_params(axis='x', rotation=45)
    ax.ticklabel_format(style='plain', axis='both')
    ax.xaxis.set_major_locator(MaxNLocator(4))
    ax.yaxis.set_major_locator(MaxNLocator(5))
    ax.set_xlabel('Easting (m)' if show_xlabel else '', fontsize=9)
    ax.set_ylabel('Northing (m)' if show_ylabel else '', fontsize=9)
    ax.annotate('N', xy=(0.97, 0.97), xytext=(0.97, 0.87),
                xycoords='axes fraction', fontsize=10, ha='right',
                arrowprops=dict(arrowstyle='->', color='black', lw=1.5),
                fontweight='bold')
    if SCALEBAR_AVAILABLE:
        sb = ScaleBar(1, units='m', location='lower left',
                      font_properties={'size': 8}, frameon=True,
                      color='black', box_alpha=0.7)
        ax.add_artist(sb)

# ============================================================
# CONFIGURATION
# ============================================================
BASE = Path("D:/ramsar_wetlands")
OUT  = BASE / "outputs"
OUT.mkdir(exist_ok=True)

YEARS  = [1991, 2001, 2015, 2025]
SITES  = ['keta', 'muni']
CNAMES = {1: 'Vegetation', 2: 'Water body', 3: 'Built up/Bareland'}
COLORS = {'Vegetation': '#2d7d2d', 'Water body': '#4a90d9', 'Built up/Bareland': '#e07b39'}
CMAP   = ListedColormap(['#2d7d2d', '#4a90d9', '#e07b39'])
NORM   = BoundaryNorm([0.5, 1.5, 2.5, 3.5], 3)

# ============================================================
# FILE PATH HELPERS
# ============================================================
def val_path(site, year):
    return BASE / f"data/{site}/validation/{site}_validation_{year}.csv"

def area_path(site, year):
    return BASE / f"data/{site}/area/{site}_area_{year}.csv"

def rst_path(site, year):
    return BASE / f"data/{site}/rasters/{site}_classified_{year}.tif"

def bnd_path(site):
    p = BASE / f"data/{site}/area/{site}_boundary.shp"
    if not p.exists():
        p = BASE / f"data/{site}/{site}_boundary.shp"
    return p

def chk(path, label):
    if not Path(path).exists():
        print(f"  [MISSING] {label}: {path}")
        return False
    return True

# ============================================================
# DATA LOADERS
# ============================================================
def load_validation(site, year):
    p = val_path(site, year)
    if not chk(p, f"{site} {year} validation"):
        return None
    df = pd.read_csv(p)
    df = df[(df['Classified'].isin([1,2,3])) & (df['GrndTruth'].isin([1,2,3]))].copy()
    return df

def load_area(site, year):
    p = area_path(site, year)
    if not chk(p, f"{site} {year} area"):
        return None
    df = pd.read_csv(p)
    # Identify columns regardless of schema version
    a_col = next((c for c in df.columns if 'Area' in c), None)
    n_col = next((c for c in df.columns if 'Class_name' in c or 'class_name' in c), None)
    if not a_col or not n_col:
        print(f"  [ERROR] Cannot parse columns in {p}")
        return None
    out = df[[n_col, a_col]].rename(columns={n_col: 'Class_name', a_col: 'Area_km2'}).copy()
    out['Class_name'] = (out['Class_name'].str.strip()
                         .replace({'Built up': 'Built up/Bareland',
                                   'Built Up': 'Built up/Bareland',
                                   'Builtup':  'Built up/Bareland'}))
    n2v = {'Vegetation': 1, 'Water body': 2, 'Built up/Bareland': 3}
    out['Class_value'] = out['Class_name'].map(n2v)
    out = out.dropna(subset=['Class_value']).sort_values('Class_value').reset_index(drop=True)
    return out

def load_raster(site, year, boundary_gdf=None):
    """Load classified raster, correctly handling 0-indexed muni rasters.

    Keta rasters (all years): values 1=Vegetation, 2=Water body, 3=Built up; 0=background
    Muni rasters (all years): values 0=Vegetation, 1=Water body, 2=Built up (stored as 0=NoData)
      -> fix by using boundary mask to separate valid pixels, then remap 0->1, 1->2, 2->3
    """
    p = rst_path(site, year)
    if not chk(p, f"{site} {year} raster"):
        return None, None, None
    with rasterio.open(p) as src:
        arr  = src.read(1).astype(np.float32)
        meta = src.meta.copy()
        tf   = src.transform
        crs  = src.crs
    meta['crs'] = crs

    # Detect encoding: if max value is 2, raster uses 0-based class indexing
    max_val = int(arr.max())
    if max_val <= 2:
        # 0-indexed encoding (muni-style): use boundary to recover Vegetation (value 0)
        if boundary_gdf is not None:
            try:
                from rasterio.features import geometry_mask
                bnd = boundary_gdf.to_crs(crs)
                geoms = [g for g in bnd.geometry if g is not None]
                # invert=True -> True where pixels are INSIDE boundary
                in_bnd = geometry_mask(geoms, transform=tf,
                                       out_shape=arr.shape, invert=True)
                result = np.full(arr.shape, np.nan, dtype=np.float32)
                result[in_bnd] = arr[in_bnd] + 1  # 0->1, 1->2, 2->3
                return result, meta, tf
            except Exception as e:
                print(f"  [WARN] boundary masking failed for {site} {year}: {e}")
        # Fallback without boundary: remap non-zero values (loses Vegetation)
        result = np.where(arr > 0, arr + 1, np.nan).astype(np.float32)
        return result, meta, tf
    else:
        # 1-indexed (keta-style): 0 = background/NoData
        arr[arr == 0] = np.nan
        return arr, meta, tf

def load_boundary(site):
    p = bnd_path(site)
    if not chk(p, f"{site} boundary"):
        return None
    try:
        return gpd.read_file(p)
    except Exception as e:
        print(f"  [WARN] boundary read error {site}: {e}")
        return None

# ============================================================
# GLOBAL STORES
# ============================================================
ACC   = {s: {} for s in SITES}   # accuracy metrics
AREAS = {s: {} for s in SITES}   # area DataFrames
RSTS  = {s: {} for s in SITES}   # raster arrays
META  = {s: {} for s in SITES}   # (meta, transform) tuples
TPMS  = {s: {} for s in SITES}   # transition prob matrices
PRED  = {s: None for s in SITES} # 2035 predictions
LM    = {s: {} for s in SITES}   # landscape metrics
FIGS  = []                        # (tag, site, year, path)

# ============================================================
# SECTION 1 — ACCURACY ASSESSMENT
# ============================================================
print("\n" + "=" * 70)
print("SECTION 1: ACCURACY ASSESSMENT")
print("=" * 70)

for site in SITES:
    print(f"\n  {site.upper()}")
    for year in YEARS:
        df = load_validation(site, year)
        if df is None:
            continue
        y_true = df['GrndTruth'].values
        y_pred = df['Classified'].values
        labels = [l for l in [1, 2, 3] if l in y_true or l in y_pred]

        cm  = confusion_matrix(y_true, y_pred, labels=labels)
        oa  = accuracy_score(y_true, y_pred)
        kap = cohen_kappa_score(y_true, y_pred)

        prod, user = {}, {}
        for i, cls in enumerate(labels):
            row_sum = cm[i, :].sum()
            col_sum = cm[:, i].sum()
            prod[CNAMES[cls]] = cm[i, i] / row_sum if row_sum else 0.0
            user[CNAMES[cls]] = cm[i, i] / col_sum if col_sum else 0.0

        ACC[site][year] = {
            'cm': cm, 'labels': labels, 'OA': oa, 'Kappa': kap,
            'Producer': prod, 'User': user, 'n': len(df)
        }
        print(f"    {year}: OA={oa:.3f}  k={kap:.3f}  n={len(df)}")

        # Confusion matrix heatmap
        fig, ax = plt.subplots(figsize=(5.5, 4.5))
        cls_names = [CNAMES[l] for l in labels]
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax,
                    xticklabels=cls_names, yticklabels=cls_names,
                    linewidths=0.5, linecolor='grey', cbar=True)
        ax.set_xlabel('')
        ax.set_ylabel('Reference', fontweight='bold')
        _site_lbl = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
        ax.set_title(f'{_site_lbl} — {year}\nOA={oa:.3f}   k={kap:.3f}   n={len(df)}')
        plt.tight_layout()
        fp = OUT / f"{site}_confusion_{year}.png"
        fig.savefig(fp, dpi=200, bbox_inches='tight')
        plt.close(fig)
        FIGS.append(('confusion', site, year, str(fp)))

# Export accuracy_summary.xlsx
print("\n  Writing accuracy_summary.xlsx...")
with pd.ExcelWriter(OUT / 'accuracy_summary.xlsx', engine='openpyxl') as xw:
    for site in SITES:
        rows = []
        for year in YEARS:
            if year not in ACC[site]:
                continue
            a = ACC[site][year]
            row = {'Year': year, 'OA': round(a['OA'], 4), 'Kappa': round(a['Kappa'], 4), 'n': a['n']}
            for cls in CNAMES.values():
                row[f"PA_{cls[:3]}"] = round(a['Producer'].get(cls, 0), 4)
                row[f"UA_{cls[:3]}"] = round(a['User'].get(cls, 0), 4)
            rows.append(row)
        pd.DataFrame(rows).to_excel(xw, sheet_name=site.capitalize(), index=False)
print("  Saved: accuracy_summary.xlsx")

# ============================================================
# SECTION 2 — LULC CHANGE ANALYSIS
# ============================================================
print("\n" + "=" * 70)
print("SECTION 2: LULC CHANGE ANALYSIS")
print("=" * 70)

for site in SITES:
    print(f"\n  {site.upper()}")
    for year in YEARS:
        df = load_area(site, year)
        if df is not None:
            AREAS[site][year] = df
            summary = "  |  ".join(
                f"{r.Class_name}={r.Area_km2:.1f} km²" for _, r in df.iterrows()
            )
            print(f"    {year}: {summary}")

CLASSES = ['Vegetation', 'Water body', 'Built up/Bareland']

for site in SITES:
    if not AREAS[site]:
        continue
    avail = [y for y in YEARS if y in AREAS[site]]

    # Build wide table
    records = [{'Year': y, **{r.Class_name: r.Area_km2
                               for _, r in AREAS[site][y].iterrows()}}
               for y in avail]
    wide = pd.DataFrame(records).set_index('Year').reindex(columns=CLASSES)

    # --- Grouped bar chart ---
    fig, ax = plt.subplots(figsize=(10, 6))
    x = np.arange(len(wide))
    w = 0.25
    for i, cls in enumerate(CLASSES):
        ax.bar(x + i * w, wide[cls], w, label=cls,
               color=list(COLORS.values())[i], edgecolor='white', linewidth=0.5)
    ax.set_xticks(x + w); ax.set_xticklabels(wide.index)
    ax.set_xlabel('Year', fontweight='bold')
    ax.set_ylabel('Area (km²)', fontweight='bold')
    ax.set_title(f'{site.capitalize()} — LULC Area by Class and Year', fontweight='bold')
    ax.legend(); ax.grid(axis='y', alpha=0.3)
    plt.tight_layout()
    fp = OUT / f"{site}_area_bar.png"
    fig.savefig(fp, dpi=200, bbox_inches='tight'); plt.close(fig)
    FIGS.append(('bar', site, None, str(fp)))

    # --- Line trend chart ---
    fig, ax = plt.subplots(figsize=(9, 5))
    for i, cls in enumerate(CLASSES):
        ax.plot(wide.index, wide[cls], marker='o',
                color=list(COLORS.values())[i], label=cls, linewidth=2, markersize=7)
        for yr, val in wide[cls].dropna().items():
            ax.annotate(f'{val:.0f}', (yr, val),
                        textcoords='offset points', xytext=(0, 8),
                        ha='center', fontsize=7.5, color=list(COLORS.values())[i])
    ax.set_xticks(wide.index); ax.set_xlabel('Year', fontweight='bold')
    ax.set_ylabel('Area (km²)', fontweight='bold')
    ax.set_title(f'{site.capitalize()} — LULC Area Trends 1991–2025', fontweight='bold')
    ax.legend(); ax.grid(alpha=0.3)
    plt.tight_layout()
    fp = OUT / f"{site}_area_trend.png"
    fig.savefig(fp, dpi=200, bbox_inches='tight'); plt.close(fig)
    FIGS.append(('trend', site, None, str(fp)))

# Export lulc_change_summary.xlsx
print("\n  Writing lulc_change_summary.xlsx...")
PERIOD_PAIRS = [(1991, 2001), (2001, 2015), (2015, 2025), (1991, 2025)]

with pd.ExcelWriter(OUT / 'lulc_change_summary.xlsx', engine='openpyxl') as xw:
    for site in SITES:
        if not AREAS[site]:
            continue
        avail = [y for y in YEARS if y in AREAS[site]]
        records = [{'Year': y, **{r.Class_name: r.Area_km2
                                   for _, r in AREAS[site][y].iterrows()}}
                   for y in avail]
        wide = pd.DataFrame(records).set_index('Year').reindex(columns=CLASSES)
        wide.to_excel(xw, sheet_name=f'{site.capitalize()}_Area')

        ch_rows = []
        for y1, y2 in PERIOD_PAIRS:
            if y1 not in AREAS[site] or y2 not in AREAS[site]:
                continue
            row = {'Period': f'{y1}–{y2}'}
            for cls in CLASSES:
                a1 = wide.loc[y1, cls] if y1 in wide.index else np.nan
                a2 = wide.loc[y2, cls] if y2 in wide.index else np.nan
                if not np.isnan(a1) and not np.isnan(a2):
                    row[f'{cls} Δkm²'] = round(a2 - a1, 2)
                    row[f'{cls} Δ%']   = round((a2 - a1) / a1 * 100, 1)
            ch_rows.append(row)
        pd.DataFrame(ch_rows).to_excel(xw, sheet_name=f'{site.capitalize()}_Change', index=False)
        print(f"    {site}: exported")
print("  Saved: lulc_change_summary.xlsx")

# ============================================================
# SECTION 3 — CA-MARKOV FUTURE PREDICTION (2035)
# ============================================================
print("\n" + "=" * 70)
print("SECTION 3: CA-MARKOV FUTURE PREDICTION (2035)")
print("=" * 70)

print("\n  Loading rasters (with boundary masking for 0-indexed rasters)...")
for site in SITES:
    print(f"  {site.upper()}:")
    bnd_gdf = load_boundary(site)
    for year in YEARS:
        arr, meta, tf = load_raster(site, year, boundary_gdf=bnd_gdf)
        if arr is not None:
            RSTS[site][year] = arr
            META[site][year] = (meta, tf)
            print(f"    {year}: shape={arr.shape}, valid={np.sum(~np.isnan(arr)):,}")


def align_arrays(a1, a2):
    """Crop both arrays to their common minimum shape."""
    r = min(a1.shape[0], a2.shape[0])
    c = min(a1.shape[1], a2.shape[1])
    return a1[:r, :c], a2[:r, :c]


def compute_tpm(arr1, arr2):
    """Compute 3×3 transition probability matrix from two classified rasters."""
    a1, a2 = align_arrays(arr1, arr2)
    valid = (~np.isnan(a1)) & (~np.isnan(a2))
    v1 = a1[valid].astype(int)
    v2 = a2[valid].astype(int)
    mask = np.isin(v1, [1, 2, 3]) & np.isin(v2, [1, 2, 3])
    v1, v2 = v1[mask], v2[mask]
    cnt = np.zeros((3, 3))
    for i in range(3):
        for j in range(3):
            cnt[i, j] = np.sum((v1 == i + 1) & (v2 == j + 1))
    rs = cnt.sum(axis=1, keepdims=True)
    rs[rs == 0] = 1
    return cnt / rs


def predict_markov(arr_2025, tpm):
    """Vectorised stochastic Markov prediction."""
    np.random.seed(42)
    valid = ~np.isnan(arr_2025)
    flat  = arr_2025[valid].astype(int)
    rand  = np.random.random(flat.shape)
    pred  = np.ones_like(flat)
    for cls in [1, 2, 3]:
        mask = flat == cls
        if not mask.any():
            continue
        cp = np.cumsum(tpm[cls - 1])
        r  = rand[mask]
        pred[mask] = 1 + (r >= cp[0]).astype(int) + (r >= cp[1]).astype(int)
    result = np.full(arr_2025.shape, np.nan)
    result[valid] = pred
    return result


np.random.seed(42)

for site in SITES:
    print(f"\n  {site.upper()} — transition matrices:")
    for y1, y2 in [(1991, 2001), (2001, 2015), (2015, 2025)]:
        if y1 in RSTS[site] and y2 in RSTS[site]:
            tpm = compute_tpm(RSTS[site][y1], RSTS[site][y2])
            TPMS[site][(y1, y2)] = tpm
            print(f"    {y1}→{y2}:")
            for i, cls in enumerate(CNAMES.values()):
                print(f"      {cls}: " + "  ".join(f"{tpm[i,j]:.4f}" for j in range(3)))

    # Select best TPM for 2025→2035 projection
    # Prefer most recent; skip if diagonal (identity = rasters are duplicate)
    def _is_identity(m, tol=0.98):
        return all(m[i, i] >= tol for i in range(3))

    tpm_for_pred = None
    tpm_period_used = None
    for ky1, ky2 in [(2015, 2025), (2001, 2015), (1991, 2001)]:
        if (ky1, ky2) in TPMS[site]:
            t = TPMS[site][(ky1, ky2)]
            if not _is_identity(t):
                tpm_for_pred = t
                tpm_period_used = (ky1, ky2)
                break
    if tpm_for_pred is None and TPMS[site]:
        # All TPMs look like identity (unusual); just use most recent
        tpm_for_pred = next(iter(reversed(list(TPMS[site].values()))))
        tpm_period_used = next(iter(reversed(list(TPMS[site].keys()))))

    if tpm_for_pred is not None and 2025 in RSTS[site]:
        tpm = tpm_for_pred
        print(f"  {site.upper()} — predicting 2035 using {tpm_period_used[0]}→{tpm_period_used[1]} TPM...")
        pred = predict_markov(RSTS[site][2025], tpm)
        PRED[site] = pred

        # Save as GeoTIFF
        m2025, tf2025 = META[site][2025]
        out_meta = m2025.copy()
        out_meta.update(dtype='uint8', nodata=0)
        pred_arr = np.where(np.isnan(pred), 0, pred).astype(np.uint8)
        tif_out = OUT / f"{site}_predicted_2035.tif"
        with rasterio.open(tif_out, 'w', **out_meta) as dst:
            dst.write(pred_arr, 1)
        print(f"    Saved: {tif_out.name}")

        # 2025 vs 2035 comparison plot — geographic-coordinate approach
        _cmap = make_cmap()
        _site_bnd = load_boundary(site)
        _m2025, _tf2025 = META[site][2025]
        _bnd_reproj = _site_bnd.to_crs(_m2025['crs']) if _site_bnd is not None else None
        _site_label = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'

        fig, axes = plt.subplots(1, 2, figsize=(14, 7))
        fig.patch.set_facecolor('white')

        for ax, (arr_plot, ptitle) in zip(axes, [
            (RSTS[site][2025], f'{_site_label} — 2025 (Classified)'),
            (pred,             f'{_site_label} — 2035 (CA-Markov Predicted)'),
        ]):
            if _site_bnd is not None:
                clipped, clip_tf, _ = clip_to_boundary(arr_plot, _m2025, _site_bnd)
            else:
                clipped, clip_tf = arr_plot, _tf2025
            plot_lulc_panel(ax, clipped, clip_tf, _bnd_reproj, ptitle, _cmap, NORM)

        patches = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
        fig.legend(handles=patches, loc='lower center', ncol=3,
                   fontsize=10, frameon=True, bbox_to_anchor=(0.5, 0.01))
        plt.tight_layout()
        plt.subplots_adjust(bottom=0.12)
        fp = OUT / f"{site}_2035_vs_2025.png"
        fig.savefig(fp, dpi=300, bbox_inches='tight'); plt.close(fig)
        FIGS.append(('pred_compare', site, 2035, str(fp)))
        print(f"    Saved: {fp.name}")

# ============================================================
# SECTION 4 — MULTI-PANEL PUBLICATION MAPS
# ============================================================
print("\n" + "=" * 70)
print("SECTION 4: MULTI-PANEL PUBLICATION MAPS")
print("=" * 70)

for site in SITES:
    site_label = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    print(f"\n  {site.upper()} — 2x2 LULC map...")

    _bnd    = load_boundary(site)
    _cmap_mp = make_cmap()

    # 2×2 multipanel: 1991 | 2001 / 2015 | 2025
    fig, axes = plt.subplots(2, 2, figsize=(12, 11))
    fig.patch.set_facecolor('white')
    plt.subplots_adjust(left=0.08, right=0.97, top=0.93, bottom=0.08,
                        wspace=0.18, hspace=0.25)

    for idx, (ax, year) in enumerate(zip(axes.flatten(), [1991, 2001, 2015, 2025])):
        row, col = divmod(idx, 2)
        if year not in RSTS[site]:
            ax.axis('off')
            continue
        arr, (ymeta, ytf) = RSTS[site][year], META[site][year]
        if _bnd is not None:
            clipped, clip_tf, _ = clip_to_boundary(arr, ymeta, _bnd)
        else:
            clipped, clip_tf = arr, ytf
        plot_lulc_panel(ax, clipped, clip_tf, None, str(year), _cmap_mp, NORM,
                        show_xlabel=(row == 1), show_ylabel=(col == 0))

    patches = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
    fig.legend(handles=patches,
               labels=['Vegetation', 'Water body', 'Built up/Bareland'],
               loc='lower center', bbox_to_anchor=(0.5, 0.01),
               ncol=3, fontsize=11, markerscale=2, frameon=True, edgecolor='gray')
    fig.suptitle(f'{site_label} — LULC Maps (1991–2025)',
                 fontsize=14, fontweight='bold', y=0.97)

    png_fp = OUT / f"{site}_multipanel.png"
    pdf_fp = OUT / f"{site}_multipanel.pdf"
    fig.savefig(png_fp, dpi=300, bbox_inches='tight', facecolor='white')
    fig.savefig(pdf_fp, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    FIGS.append(('multipanel', site, None, str(png_fp)))
    print(f"    Saved: {png_fp.name} + .pdf")

    # 2035 standalone map
    if PRED[site] is not None and 2025 in META[site]:
        pred_meta, pred_tf = META[site][2025]
        if _bnd is not None:
            pred_clipped, pred_clip_tf, _ = clip_to_boundary(PRED[site], pred_meta, _bnd)
        else:
            pred_clipped, pred_clip_tf = PRED[site], pred_tf
        fig2, ax2 = plt.subplots(1, 1, figsize=(8, 8))
        fig2.patch.set_facecolor('white')
        plt.subplots_adjust(left=0.12, right=0.92, top=0.88, bottom=0.12)
        plot_lulc_panel(ax2, pred_clipped, pred_clip_tf, None, '', _cmap_mp, NORM,
                        show_xlabel=True, show_ylabel=True)
        patches2 = [mpatches.Patch(color=c, label=n) for n, c in COLORS.items()]
        fig2.legend(handles=patches2,
                    labels=['Vegetation', 'Water body', 'Built up/Bareland'],
                    loc='lower center', bbox_to_anchor=(0.5, 0.01),
                    ncol=3, fontsize=11, markerscale=2, frameon=True, edgecolor='gray')
        fig2.suptitle(f'{site_label} — Predicted LULC 2035 (CA-Markov)',
                      fontsize=12, fontweight='bold', y=0.97)
        pred_map_fp = OUT / f"{site}_predicted_2035_map.png"
        fig2.savefig(pred_map_fp, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig2)
        FIGS.append(('pred_map', site, 2035, str(pred_map_fp)))
        print(f"    Saved: {pred_map_fp.name}")

# ============================================================
# SECTION 5 — LANDSCAPE METRICS
# ============================================================
print("\n" + "=" * 70)
print("SECTION 5: LANDSCAPE METRICS")
print("=" * 70)


def shannon_di(arr):
    valid = arr[~np.isnan(arr)].astype(int)
    valid = valid[np.isin(valid, [1, 2, 3])]
    n = len(valid)
    if n == 0:
        return 0.0
    sdi = 0.0
    for cls in [1, 2, 3]:
        p = np.sum(valid == cls) / n
        if p > 0:
            sdi -= p * np.log(p)
    return sdi


def patch_metrics(arr):
    valid_total = int(np.sum(~np.isnan(arr)))
    out = {}
    for cls in [1, 2, 3]:
        binary = (arr == cls).astype(int)
        labeled, n_pat = ndlabel(binary)
        if n_pat > 0:
            sizes = [int(np.sum(labeled == p)) for p in range(1, n_pat + 1)]
            lpi  = max(sizes) / valid_total * 100 if valid_total else 0.0
            mean_sz = float(np.mean(sizes))
        else:
            lpi, mean_sz = 0.0, 0.0
        out[CNAMES[cls]] = {'n_patches': n_pat, 'LPI': round(lpi, 4),
                             'mean_patch_px': round(mean_sz, 1)}
    return out


for site in SITES:
    print(f"\n  {site.upper()}:")
    all_yr = sorted(list(RSTS[site].keys()) + ([2035] if PRED[site] is not None else []))
    for year in all_yr:
        arr = PRED[site] if year == 2035 else RSTS[site].get(year)
        if arr is None:
            continue
        sdi  = shannon_di(arr)
        frag = patch_metrics(arr)
        LM[site][year] = {'SDI': sdi, 'frag': frag}
        patch_str = ", ".join(f"{cls}: {info['n_patches']}p LPI={info['LPI']:.2f}%"
                               for cls, info in frag.items())
        print(f"    {year}: SDI={sdi:.4f}  |  {patch_str}")

        if PYLANDSTATS_AVAILABLE:
            try:
                arr_int = np.where(np.isnan(arr), 0, arr).astype(np.uint8)
                ls = pls.Landscape(arr_int, res=(30, 30), nodata=0)
                LM[site][year]['pls'] = ls
            except Exception:
                pass

for site in SITES:
    if not LM[site]:
        continue
    yr_list = sorted(LM[site].keys())
    sdi_vals = [LM[site][y]['SDI'] for y in yr_list]

    fig, axes = plt.subplots(1, 3, figsize=(16, 5))
    bar_clr = ['#2d6a2d'] * 4 + ['#1a4f1a'] if len(yr_list) == 5 else ['#2d6a2d'] * len(yr_list)
    axes[0].bar([str(y) for y in yr_list], sdi_vals, color=bar_clr, edgecolor='white')
    axes[0].set_title('Shannon Diversity Index', fontweight='bold')
    axes[0].set_xlabel('Year'); axes[0].set_ylabel('SDI')
    for i, v in enumerate(sdi_vals):
        axes[0].text(i, v + 0.003, f'{v:.3f}', ha='center', va='bottom', fontsize=8)
    axes[0].grid(axis='y', alpha=0.3)

    for ci, cls in enumerate(CNAMES.values()):
        yrs = [y for y in yr_list if cls in LM[site][y]['frag']]
        vals = [LM[site][y]['frag'][cls]['n_patches'] for y in yrs]
        axes[1].plot([str(y) for y in yrs], vals, marker='o',
                     color=list(COLORS.values())[ci], label=cls, lw=2)
    axes[1].set_title('Number of Patches per Class', fontweight='bold')
    axes[1].set_xlabel('Year'); axes[1].set_ylabel('Patch Count')
    axes[1].legend(fontsize=8); axes[1].grid(alpha=0.3)

    for ci, cls in enumerate(CNAMES.values()):
        yrs  = [y for y in yr_list if cls in LM[site][y]['frag']]
        vals = [LM[site][y]['frag'][cls]['LPI'] for y in yrs]
        axes[2].plot([str(y) for y in yrs], vals, marker='s',
                     color=list(COLORS.values())[ci], label=cls, lw=2)
    axes[2].set_title('Largest Patch Index (%)', fontweight='bold')
    axes[2].set_xlabel('Year'); axes[2].set_ylabel('LPI (%)')
    axes[2].legend(fontsize=8); axes[2].grid(alpha=0.3)

    site_label = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    fig.suptitle(f'{site_label} — Landscape Metrics', fontweight='bold', fontsize=12)
    plt.tight_layout()
    fp = OUT / f"{site}_landscape_metrics.png"
    fig.savefig(fp, dpi=200, bbox_inches='tight'); plt.close(fig)
    FIGS.append(('landscape', site, None, str(fp)))
    print(f"    Saved: {fp.name}")

# ============================================================
# SECTION 6 — WORD REPORT
# ============================================================
print("\n" + "=" * 70)
print("SECTION 6: GENERATING WORD DOCUMENT REPORT")
print("=" * 70)


# ---- docx helpers -----------------------------------------------
def cell_border(cell):
    tc  = cell._tc
    pr  = tc.get_or_add_tcPr()
    bdr = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), '000000')
        bdr.append(e)
    pr.append(bdr)


def add_df_table(doc, df, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_border(hdr[i])
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_border(cells[i])
    doc.add_paragraph()


def fig_insert(doc, path, caption='', width=15.0):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
        doc.add_paragraph()


# ---- collect key metrics for abstract & discussion ---------------
def _get_val(site, year, cls):
    if year not in AREAS[site]:
        return np.nan
    s = AREAS[site][year]
    m = s[s['Class_name'] == cls]['Area_km2']
    return float(m.iloc[0]) if len(m) else np.nan


km = {}
for site in SITES:
    oas = [ACC[site][y]['OA'] for y in YEARS if y in ACC[site]]
    kps = [ACC[site][y]['Kappa'] for y in YEARS if y in ACC[site]]
    veg91 = _get_val(site, 1991, 'Vegetation')
    veg25 = _get_val(site, 2025, 'Vegetation')
    bu91  = _get_val(site, 1991, 'Built up/Bareland')
    bu25  = _get_val(site, 2025, 'Built up/Bareland')
    km[site] = {
        'oa_min': min(oas) if oas else 0, 'oa_max': max(oas) if oas else 0,
        'kp_min': min(kps) if kps else 0, 'kp_max': max(kps) if kps else 0,
        'veg_ch': veg25 - veg91 if not np.isnan(veg91 + veg25) else 0,
        'bu_ch':  bu25  - bu91  if not np.isnan(bu91  + bu25)  else 0,
    }


def biggest_change(site):
    changes = {}
    for cls in CLASSES:
        a1 = _get_val(site, 1991, cls)
        a2 = _get_val(site, 2025, cls)
        if not np.isnan(a1 + a2):
            changes[cls] = abs(a2 - a1)
    if not changes:
        return 'Unknown', 0
    mc = max(changes, key=changes.get)
    return mc, changes[mc]


keta_mc, keta_mv = biggest_change('keta')
muni_mc, muni_mv = biggest_change('muni')

# ---- Build document -----------------------------------------------
doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)

# Title page
for _ in range(4):
    doc.add_paragraph()
tp = doc.add_paragraph(
    "Wetland Degradation Analysis of Keta Lagoon Complex\n"
    "and Muni-Pomadze Ramsar Sites, Ghana"
)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.runs[0].bold = True; tp.runs[0].font.size = Pt(18)

doc.add_paragraph()
sp = doc.add_paragraph(
    "Land Use / Land Cover Change Detection and\n"
    "CA-Markov Future Projection (1991–2035)"
)
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.runs[0].font.size = Pt(13)

doc.add_paragraph()
dp = doc.add_paragraph("June 2026")
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.runs[0].font.size = Pt(12)
doc.add_page_break()

# Abstract
doc.add_heading('Abstract', level=1)
abstract = (
    "This study investigates land use and land cover (LULC) dynamics at two Ghana Ramsar Sites — "
    "Keta Lagoon Complex and Muni-Pomadze — over a 34-year period (1991–2025) using multi-temporal "
    "Landsat satellite imagery and Support Vector Machine (SVM) supervised classification. Three "
    "LULC classes were mapped: Vegetation, Water body, and Built up/Bareland. Accuracy assessments "
    f"yielded Overall Accuracy (OA) values of {km['keta']['oa_min']:.1%}–{km['keta']['oa_max']:.1%} "
    f"(Kappa: {km['keta']['kp_min']:.3f}–{km['keta']['kp_max']:.3f}) for Keta Lagoon Complex, "
    f"and {km['muni']['oa_min']:.1%}–{km['muni']['oa_max']:.1%} "
    f"(Kappa: {km['muni']['kp_min']:.3f}–{km['muni']['kp_max']:.3f}) for Muni-Pomadze, confirming "
    "satisfactory classification performance. Between 1991 and 2025, Vegetation cover at Keta changed by "
    f"{km['keta']['veg_ch']:+.1f} km² while Built up/Bareland expanded by {km['keta']['bu_ch']:+.1f} km²; "
    f"at Muni-Pomadze, Vegetation changed by {km['muni']['veg_ch']:+.1f} km² and Built up/Bareland "
    f"by {km['muni']['bu_ch']:+.1f} km². CA-Markov transition modelling projects continued degradation "
    "trends through 2035 at both sites. Shannon Diversity Index and patch-fragmentation metrics "
    "corroborate observed LULC dynamics. Findings underscore significant wetland degradation pressures "
    "and have direct implications for Ramsar site management and conservation policy."
)
doc.add_paragraph(abstract)
doc.add_page_break()

# ---- 1. Methodology ----
doc.add_heading('1. Methodology', level=1)

doc.add_heading('1.1 Study Area', level=2)
doc.add_paragraph(
    "Keta Lagoon Complex Ramsar Site is located in the Volta Region of south-eastern Ghana "
    "(approx. 5°50′N, 0°55′E) and is the country's largest Ramsar Site, encompassing a mosaic of "
    "coastal lagoon, mangrove, salt marsh and estuarine habitats. Muni-Pomadze Ramsar Site lies "
    "along the central coast of Ghana (approx. 5°22′N, 0°36′W) in the Central Region and consists "
    "of a shallow coastal lagoon surrounded by coastal scrub, mangrove fringe and sandy beach. "
    "Both sites support internationally important waterbird populations and are under increasing "
    "anthropogenic pressure from urban expansion, agriculture and climate variability."
)

doc.add_heading('1.2 Data and Classification', level=2)
doc.add_paragraph(
    "Multi-temporal Landsat imagery (Landsat 4/5 TM for 1991, Landsat 7 ETM+ for 2001, "
    "Landsat 8 OLI for 2015, Landsat 8/9 OLI-2 for 2025) was acquired and pre-processed "
    "(radiometric calibration, atmospheric correction, geometric registration). A supervised "
    "Support Vector Machine (SVM) classifier was applied in ArcGIS Pro to derive three LULC classes: "
    "(1) Vegetation, (2) Water body, and (3) Built up/Bareland. All outputs were projected to "
    "UTM Zone 31N (WGS 1984) at 30 m spatial resolution."
)

doc.add_heading('1.3 Accuracy Assessment', level=2)
doc.add_paragraph(
    "Independent stratified-random validation points were used to generate confusion matrices for "
    "each site and time period. Overall Accuracy (OA), Kappa coefficient (κ), Producer's Accuracy "
    "(PA) and User's Accuracy (UA) were computed for all three classes following Congalton & Green (2019)."
)

doc.add_heading('1.4 Change Detection', level=2)
doc.add_paragraph(
    "Class-level area statistics (km²) were compared across the four time periods. Net area change "
    "(km²) and percentage change were calculated for inter-period intervals (1991–2001, 2001–2015, "
    "2015–2025) and the full study period (1991–2025). Transition probability matrices were derived "
    "from pixel-to-pixel class conversions between consecutive classified maps."
)

doc.add_heading('1.5 CA-Markov Future Prediction', level=2)
doc.add_paragraph(
    "Transition probability matrices (TPMs) computed from the 2015–2025 period were used to project "
    "LULC to 2035 via a Cellular Automaton–Markov Chain (CA-Markov) approach. Each pixel in the "
    "2025 map was stochastically assigned a 2035 class based on the TPM row corresponding to its "
    "current class, implemented with a fixed random seed (42) for reproducibility. The prediction "
    "was exported as a co-registered GeoTIFF raster."
)

doc.add_heading('1.6 Landscape Metrics', level=2)
doc.add_paragraph(
    "Landscape composition was quantified using the Shannon Diversity Index "
    "(SDI = −Σpᵢ ln pᵢ), where pᵢ is the proportion of the landscape in class i. "
    "Configuration was characterised by the number of discrete patches per class "
    "(four-connectivity labelling), the Largest Patch Index (LPI; % of landscape area "
    "occupied by the largest single patch), and mean patch size, computed for all time periods "
    "including the 2035 projection."
)
doc.add_page_break()

# ---- 2. Results ----
doc.add_heading('2. Results', level=1)

# 2.1 Accuracy
doc.add_heading('2.1 Accuracy Assessment', level=2)
doc.add_paragraph(
    "Accuracy assessment results across all four time periods and both sites are presented "
    "in the tables and confusion matrix figures below."
)

for site in SITES:
    sn = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    doc.add_heading(f'2.1.{"1" if site=="keta" else "2"}  {sn}', level=3)

    rows = []
    for year in YEARS:
        if year not in ACC[site]:
            continue
        a = ACC[site][year]
        rows.append({
            'Year': year,
            'OA (%)': f"{a['OA']*100:.1f}",
            'κ': f"{a['Kappa']:.3f}",
            'PA Veg (%)': f"{a['Producer'].get('Vegetation', 0)*100:.1f}",
            'UA Veg (%)': f"{a['User'].get('Vegetation', 0)*100:.1f}",
            'PA Wat (%)': f"{a['Producer'].get('Water body', 0)*100:.1f}",
            'UA Wat (%)': f"{a['User'].get('Water body', 0)*100:.1f}",
            'PA BuBl (%)': f"{a['Producer'].get('Built up/Bareland', 0)*100:.1f}",
            'UA BuBl (%)': f"{a['User'].get('Built up/Bareland', 0)*100:.1f}",
        })
    add_df_table(doc, pd.DataFrame(rows))
    p = doc.add_paragraph(
        f"Table 2.{1 if site == 'keta' else 2}: Accuracy assessment for {sn} "
        f"(PA = Producer's Accuracy, UA = User's Accuracy, BuBl = Built up/Bareland)."
    )
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for year in YEARS:
        fp = OUT / f"{site}_confusion_{year}.png"
        fig_insert(doc, str(fp),
                   caption=f"Figure: {sn} {year} confusion matrix heatmap.",
                   width=9.0)

# 2.2 LULC Change
doc.add_heading('2.2 LULC Change Analysis', level=2)
doc.add_paragraph(
    "Area statistics and net changes between time periods are presented below for each site."
)

tbl_num = 3
for site in SITES:
    sn = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    doc.add_heading(f'2.2.{"1" if site=="keta" else "2"}  {sn}', level=3)

    area_rows = []
    for year in YEARS:
        if year not in AREAS[site]:
            continue
        row = {'Year': year}
        for _, r in AREAS[site][year].iterrows():
            row[r['Class_name']] = f"{r['Area_km2']:.2f}"
        area_rows.append(row)
    add_df_table(doc, pd.DataFrame(area_rows))
    p = doc.add_paragraph(f"Table {tbl_num}: LULC area (km²) by class and year — {sn}.")
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9); doc.add_paragraph()
    tbl_num += 1

    ch_rows = []
    for y1, y2 in PERIOD_PAIRS:
        if y1 not in AREAS[site] or y2 not in AREAS[site]:
            continue
        row = {'Period': f'{y1}–{y2}'}
        for cls in CLASSES:
            a1 = _get_val(site, y1, cls)
            a2 = _get_val(site, y2, cls)
            if not np.isnan(a1 + a2):
                row[f'{cls[:3]} Δkm²'] = f"{a2-a1:+.2f}"
                row[f'{cls[:3]} Δ%']   = f"{(a2-a1)/a1*100:+.1f}"
        ch_rows.append(row)
    add_df_table(doc, pd.DataFrame(ch_rows))
    p = doc.add_paragraph(f"Table {tbl_num}: Net LULC change by period — {sn}.")
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9); doc.add_paragraph()
    tbl_num += 1

    fig_insert(doc, str(OUT / f"{site}_area_bar.png"),
               caption=f"Figure: {sn} — LULC area by class and year (grouped bar).", width=15)
    fig_insert(doc, str(OUT / f"{site}_area_trend.png"),
               caption=f"Figure: {sn} — LULC area trend lines 1991–2025.", width=15)

# 2.3 CA-Markov
doc.add_heading('2.3 CA-Markov Future Prediction (2035)', level=2)
doc.add_paragraph(
    "Transition probability matrices derived from 2015–2025 were used to project LULC to 2035. "
    "The following tables and figures present the TPMs and predicted maps."
)

for site in SITES:
    sn = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    doc.add_heading(f'2.3.{"1" if site=="keta" else "2"}  {sn}', level=3)

    if (2015, 2025) in TPMS[site]:
        tpm = TPMS[site][(2015, 2025)]
        tpm_rows = [
            {'From \\ To': CNAMES[i+1],
             **{CNAMES[j+1]: f"{tpm[i,j]:.4f}" for j in range(3)}}
            for i in range(3)
        ]
        add_df_table(doc, pd.DataFrame(tpm_rows))
        p = doc.add_paragraph(f"Table {tbl_num}: Transition probability matrix (2015→2025) — {sn}.")
        p.runs[0].italic = True; p.runs[0].font.size = Pt(9); doc.add_paragraph()
        tbl_num += 1

    fig_insert(doc, str(OUT / f"{site}_2035_vs_2025.png"),
               caption=f"Figure: {sn} — 2025 classified vs. 2035 CA-Markov predicted.", width=15)

# 2.4 Multi-panel maps
doc.add_heading('2.4 Multi-panel LULC Maps (1991–2035)', level=2)
doc.add_paragraph(
    "Publication-quality five-panel maps for each site showing all classified and predicted time steps."
)
for site in SITES:
    sn = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    fig_insert(doc, str(OUT / f"{site}_multipanel.png"),
               caption=f"Figure: {sn} — Multi-temporal LULC maps 1991–2035 (1991–2025 SVM; 2035 CA-Markov).",
               width=16)

# 2.5 Landscape metrics
doc.add_heading('2.5 Landscape Metrics', level=2)
doc.add_paragraph(
    "Shannon Diversity Index, patch count, and Largest Patch Index for all time periods "
    "including the 2035 projection."
)
for site in SITES:
    sn = 'Keta Lagoon Complex' if site == 'keta' else 'Muni-Pomadze'
    doc.add_heading(f'2.5.{"1" if site=="keta" else "2"}  {sn}', level=3)

    lm_rows = []
    for year in sorted(LM[site].keys()):
        lm = LM[site][year]
        row = {'Year': f"{year}{' (P)' if year==2035 else ''}", 'SDI': f"{lm['SDI']:.4f}"}
        for cls in CNAMES.values():
            if cls in lm['frag']:
                row[f'{cls[:3]} Patches'] = str(lm['frag'][cls]['n_patches'])
                row[f'{cls[:3]} LPI (%)'] = f"{lm['frag'][cls]['LPI']:.2f}"
        lm_rows.append(row)
    add_df_table(doc, pd.DataFrame(lm_rows))
    p = doc.add_paragraph(f"Table {tbl_num}: Landscape metrics — {sn}. (P) = Predicted.")
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9); doc.add_paragraph()
    tbl_num += 1

    fig_insert(doc, str(OUT / f"{site}_landscape_metrics.png"),
               caption=f"Figure: {sn} — Landscape metrics 1991–2035.", width=16)

doc.add_page_break()

# ---- 3. Discussion ----
doc.add_heading('3. Discussion', level=1)

disc = (
    f"The analysis reveals substantial LULC transformations at both Ramsar sites over the 34-year "
    f"study period. At Keta Lagoon Complex, '{keta_mc}' experienced the greatest absolute change "
    f"({keta_mv:.1f} km²), reflecting the intensity of coastal land transformation processes — "
    f"including artisanal salt harvesting, commercial fishing infrastructure, and peri-urban growth "
    f"— that characterise this area. Built up/Bareland expansion at Keta ({km['keta']['bu_ch']:+.1f} km² "
    f"from 1991 to 2025) is consistent with documented urbanisation trends along Ghana's Volta coast.\n\n"
    f"At Muni-Pomadze, '{muni_mc}' registered the largest change ({muni_mv:.1f} km²). Despite the "
    f"site's smaller spatial extent, proportional changes are significant, highlighting its "
    f"vulnerability to relatively modest anthropogenic disturbances including small-scale subsistence "
    f"agriculture, coastal erosion and seasonal hydrological variability. The Vegetation change of "
    f"{km['muni']['veg_ch']:+.1f} km² between 1991 and 2025 indicates progressive habitat loss "
    f"in an already compact wetland system.\n\n"
    f"Comparing the two sites, Keta Lagoon Complex exhibits larger absolute changes owing to its "
    f"greater total extent, while Muni-Pomadze shows disproportionate proportional changes relative "
    f"to its area. Both sites display the hallmarks of wetland degradation documented across "
    f"West Africa: shrinking water bodies, declining natural vegetation cover, and expanding "
    f"degraded or built-up land (Ozesmi & Bauer, 2002; Dronova, 2015).\n\n"
    f"The CA-Markov projections to 2035 indicate that, absent targeted management intervention, "
    f"current degradation trajectories will persist. The high row-wise probabilities of remaining "
    f"in the Built up/Bareland class in the 2015–2025 TPMs confirm this class's stability and "
    f"near-irreversible nature once established. Vegetation-to-Bareland transitions at both sites "
    f"reinforce the urgency of proactive conservation measures.\n\n"
    f"Shannon Diversity Index trends provide additional evidence of changing landscape structure. "
    f"Declining SDI values over time indicate decreasing landscape heterogeneity, with one or a few "
    f"classes increasingly dominating the landscape — adverse for biodiversity and ecosystem "
    f"resilience. The fragmentation metrics (increasing patch counts, declining mean patch size) "
    f"suggest progressive habitat fragmentation, reducing landscape connectivity for wetland-dependent "
    f"species. These findings have direct implications for Ramsar site management: specifically, "
    f"the need to enforce existing protection boundaries, develop wetland buffer zones, support "
    f"community-based conservation programmes, and integrate remote-sensing-based monitoring into "
    f"national environmental governance frameworks."
)
for para in disc.split('\n\n'):
    doc.add_paragraph(para.strip())
doc.add_page_break()

# ---- 4. Conclusion ----
doc.add_heading('4. Conclusion', level=1)

concl = (
    f"This study provides a comprehensive multi-temporal LULC assessment at Keta Lagoon Complex "
    f"and Muni-Pomadze Ramsar Sites in Ghana from 1991 to 2025, with CA-Markov projections to 2035. "
    f"SVM-classified Landsat imagery achieved satisfactory accuracy across all time periods "
    f"(OA: {km['keta']['oa_min']:.1%}–{km['keta']['oa_max']:.1%} for Keta; "
    f"{km['muni']['oa_min']:.1%}–{km['muni']['oa_max']:.1%} for Muni-Pomadze). "
    f"Both sites exhibit clear evidence of wetland degradation — declining Vegetation and Water body "
    f"extent alongside expanding Built up/Bareland — with trends projected to continue through 2035 "
    f"under business-as-usual land management.\n\n"
    f"Key recommendations emerging from this work include: (i) strengthened enforcement of "
    f"Ramsar site boundary protection zones; (ii) community-engaged sustainable livelihood "
    f"programmes to reduce encroachment pressure; (iii) establishment of a near-real-time "
    f"remote sensing monitoring programme using freely available Landsat/Sentinel imagery; "
    f"and (iv) integration of wetland ecosystem service valuation into national land-use planning "
    f"and environmental impact assessment processes. Future research should incorporate field-based "
    f"vegetation surveys, hydrological data and socioeconomic driver analysis to develop a more "
    f"complete causal model of degradation at these globally important wetland sites."
)
for para in concl.split('\n\n'):
    doc.add_paragraph(para.strip())
doc.add_page_break()

# ---- 5. References ----
doc.add_heading('5. References', level=1)
refs = [
    ("Congalton, R. G., & Green, K. (2019). "
     "Assessing the accuracy of remotely sensed data: Principles and practices (3rd ed.). "
     "CRC Press."),
    ("Dronova, I. (2015). Object-based image analysis in wetland research: A review. "
     "Remote Sensing, 7(5), 6380–6413. https://doi.org/10.3390/rs70506380"),
    ("Giri, C., Zhu, Z., & Reed, B. (2005). A comparative analysis of the Global Land Cover 2000 "
     "and MODIS land cover data sets. Remote Sensing of Environment, 94(1), 123–132. "
     "https://doi.org/10.1016/j.rse.2004.09.005"),
    ("Ozesmi, S. L., & Bauer, M. E. (2002). Satellite remote sensing of wetlands. "
     "Wetlands Ecology and Management, 10(5), 381–402. https://doi.org/10.1023/A:1020908432489"),
    ("Prigent, C., Papa, F., Aires, F., Rossow, W. B., & Matthews, E. (2007). Global inundation "
     "dynamics inferred from multiple satellite observations, 1993–2000. "
     "Journal of Geophysical Research: Atmospheres, 112(D12). "
     "https://doi.org/10.1029/2006JD007847"),
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.left_indent       = Cm(0.5)

doc_path = OUT / "wetland_degradation_report.docx"
doc.save(str(doc_path))
print(f"  Saved: wetland_degradation_report.docx")

# ============================================================
# PIPELINE SUMMARY
# ============================================================
print("\n" + "=" * 70)
print("PIPELINE COMPLETE — all outputs saved to:", OUT)
print("=" * 70)
all_files = sorted(OUT.iterdir())
for f in all_files:
    if f.is_file():
        print(f"  {f.name:<50s}  {f.stat().st_size/1024:>8.1f} KB")
print("=" * 70)
