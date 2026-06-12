#!/usr/bin/env python3
"""
fix_docx_final.py — Apply all 8 docx fixes and save as
wetland_degradation_report_FINAL.docx
"""
import warnings; warnings.filterwarnings('ignore')
import shutil
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

BASE     = Path("D:/ramsar_wetlands")
OUT      = BASE / "outputs"
SRC_DOCX = OUT / "wetland_degradation_report.docx"
DST_DOCX = OUT / "wetland_degradation_report_FINAL.docx"

shutil.copy2(SRC_DOCX, DST_DOCX)
doc = Document(str(DST_DOCX))

print("=" * 70)
print("FINAL DOCX FIXES")
print("=" * 70)

# ── helpers ────────────────────────────────────────────────────────────────
def para_text(p):
    return ''.join(r.text for r in p.runs)

def find_para(*keywords, start=0):
    """First paragraph containing ALL keywords."""
    for i, p in enumerate(doc.paragraphs):
        if i < start: continue
        t = p.text
        if all(k in t for k in keywords):
            return i, p
    return -1, None

def append_sentence_to_para(para, sentence):
    """Add a sentence (space + text) to an existing paragraph as a new run."""
    run = para.add_run(' ' + sentence)
    # Inherit font size from first run if available
    if len(para.runs) > 1 and para.runs[0].font.size:
        run.font.size = para.runs[0].font.size

def replace_in_para(para, old, new):
    """Replace text across all runs in a paragraph (handles cross-run spans)."""
    full = para.text
    if old not in full:
        return False
    # Try single-run replacement first
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Multi-run: rebuild first run, clear others (rare but safe fallback)
    new_full = full.replace(old, new)
    for i, run in enumerate(para.runs):
        run.text = new_full if i == 0 else ''
    return True

def insert_para_after(anchor_para, text, italic=False, size_pt=None,
                      alignment=None, hanging=False):
    """Insert a new paragraph immediately after anchor_para."""
    new_p = OxmlElement('w:p')
    anchor_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np_obj = Paragraph(new_p, doc)
    run = np_obj.add_run(text)
    if italic:       run.italic = True
    if size_pt:      run.font.size = Pt(size_pt)
    if alignment:    np_obj.alignment = alignment
    if hanging:
        np_obj.paragraph_format.first_line_indent = Cm(-0.5)
        np_obj.paragraph_format.left_indent       = Cm(0.5)
    return np_obj

def insert_para_before(anchor_para, text, **kwargs):
    """Insert a new paragraph immediately before anchor_para."""
    prev = anchor_para._p.getprevious()
    if prev is not None:
        dummy = OxmlElement('w:p')
        prev.addnext(dummy)
        anchor_para._p.addprevious(dummy)
        from docx.text.paragraph import Paragraph
        np_obj = Paragraph(dummy, doc)
        run = np_obj.add_run(text)
        return np_obj
    # Fallback: just insert after anchor
    return insert_para_after(anchor_para, text, **kwargs)

def cell_border(cell):
    tc  = cell._tc
    pr  = tc.get_or_add_tcPr()
    bdr = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), '000000')
        bdr.append(e)
    pr.append(bdr)

# ── LOAD AREA DATA FOR FIX 6 ────────────────────────────────────────────
YEARS = [1991, 2001, 2015, 2025]
CLASSES = ['Vegetation', 'Water body', 'Built up/Bareland']
AREA_TOTALS = {}

for site in ['keta', 'muni']:
    AREA_TOTALS[site] = {}
    rows_by_year = {}
    for year in YEARS:
        p = BASE / f"data/{site}/area/{site}_area_{year}.csv"
        df = pd.read_csv(p)
        a_col = next(c for c in df.columns if 'Area' in c)
        n_col = next(c for c in df.columns if 'Class_name' in c or 'class_name' in c)
        df2   = df[[n_col, a_col]].rename(columns={n_col: 'Class', a_col: 'Area'})
        df2['Class'] = (df2['Class'].str.strip()
                        .replace({'Built up': 'Built up/Bareland', 'Built Up': 'Built up/Bareland',
                                  'Builtup': 'Built up/Bareland'}))
        rows_by_year[year] = {r['Class']: r['Area'] for _, r in df2.iterrows()}
    # Column sums (across years) for each class
    for cls in CLASSES:
        AREA_TOTALS[site][cls] = round(sum(rows_by_year[y].get(cls, 0) for y in YEARS), 2)
    # Per-year totals
    AREA_TOTALS[site]['per_year'] = {
        y: round(sum(rows_by_year[y].get(cls, 0) for cls in CLASSES), 2) for y in YEARS
    }
    print(f"  {site} column sums: {AREA_TOTALS[site]}")

print()

# ════════════════════════════════════════════════════════════════════════════
# FIX 1 — Discussion: add sentence after "...proactive conservation measures."
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 1] Adding sentence after 'proactive conservation measures.'...")
idx, p = find_para("reinforce the urgency of proactive conservation measures")
if p is not None:
    target_end = "reinforce the urgency of proactive conservation measures."
    new_sent   = ("The relatively high Built up/Bareland-to-Vegetation transition "
                  "probability at Keta (0.53) likely reflects seasonal dynamics in "
                  "salt marsh and tidal flat areas rather than permanent land recovery, "
                  "and should be interpreted with caution.")
    if new_sent[:30] not in p.text:
        # Append to the paragraph if target sentence is at end, else insert after
        if p.text.rstrip().endswith(target_end.rstrip('.')):
            append_sentence_to_para(p, new_sent)
            print(f"  Appended to para {idx}.")
        else:
            ok = replace_in_para(p, target_end, target_end + ' ' + new_sent)
            print(f"  Inserted inline: {ok}")
    else:
        print("  Already present — skipped.")
else:
    print("  [WARN] Target paragraph not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 2 — Discussion: add sentence after "...landscape connectivity..."
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 2] Adding sentence after 'reducing landscape connectivity...'...")
TARGET_F2 = "reducing landscape connectivity for wetland-dependent species."
NEW_F2    = ("The projected increase in Vegetation LPI at Keta by 2035 reflects "
             "spatial consolidation of remaining vegetated areas rather than net gain, "
             "as total Vegetation extent continues to decline under the CA-Markov projection.")
idx, p = find_para("reducing landscape connectivity for wetland-dependent species")
if p is not None:
    if NEW_F2[:30] not in p.text:
        ok = replace_in_para(p, TARGET_F2, TARGET_F2 + ' ' + NEW_F2)
        print(f"  Inserted inline: {ok}")
    else:
        print("  Already present — skipped.")
else:
    print("  [WARN] Target paragraph not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 3 — Section 1.3: add Pontius & Millones disclaimer after Foody citation
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 3] Adding Pontius & Millones sentence in Section 1.3...")
TARGET_F3 = "following Foody (2002) and Congalton & Green (2019)."
NEW_F3    = ("While Pontius & Millones (2011) question the utility of the Kappa "
             "coefficient, it is retained here for comparability with existing "
             "wetland remote sensing literature.")
idx, p = find_para("following Foody (2002) and Congalton & Green (2019)")
if p is not None:
    if NEW_F3[:30] not in p.text:
        ok = replace_in_para(p, TARGET_F3, TARGET_F3 + ' ' + NEW_F3)
        print(f"  Inserted inline: {ok}")
    else:
        print("  Already present — skipped.")
else:
    # Try without Foody citation (pre-fix version)
    TARGET_F3b = "following Congalton & Green (2019)."
    idx, p = find_para("following Congalton & Green (2019)")
    if p is not None and NEW_F3[:30] not in p.text:
        ok = replace_in_para(p, TARGET_F3b, TARGET_F3b + ' ' + NEW_F3)
        print(f"  Inserted via fallback: {ok}")
    else:
        print("  [WARN] Section 1.3 accuracy paragraph not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 4 — Section 1.2: add Vapnik (1995) citation
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 4] Adding Vapnik (1995) citation in Section 1.2...")
OLD_F4 = "A supervised Support Vector Machine (SVM) classifier was applied in ArcGIS Pro"
NEW_F4 = "A supervised Support Vector Machine (SVM) classifier (Vapnik, 1995) was applied in ArcGIS Pro"
idx, p = find_para("Support Vector Machine (SVM) classifier")
if p is not None:
    if "(Vapnik, 1995)" not in p.text:
        ok = replace_in_para(p, OLD_F4, NEW_F4)
        print(f"  Replaced: {ok}")
    else:
        print("  Already present — skipped.")
else:
    print("  [WARN] SVM paragraph not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 5 — Discussion para 1: add Schiavina et al. (2022) at end
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 5] Adding Schiavina et al. (2022) to Discussion para 1...")
SCHIAVINA_SENT = ("Population growth along Ghana's coast (Schiavina et al., 2022) "
                  "is a primary driver of encroachment pressure at both Ramsar sites.")
idx, p = find_para("The analysis reveals substantial LULC transformations")
if p is not None:
    if "Schiavina" not in p.text:
        append_sentence_to_para(p, SCHIAVINA_SENT)
        print(f"  Appended to Discussion para 1 (para {idx}).")
    else:
        print("  Already present — skipped.")
else:
    print("  [WARN] Discussion first paragraph not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 6 — Add Total rows + footnotes to Tables 3 and 5
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 6] Adding Total rows and footnotes to Tables 3 and 5...")

def add_total_row_and_footnote(tbl_idx, site, caption_keywords):
    tbl = doc.tables[tbl_idx]
    # Check if Total row already present
    if any(row.cells[0].text.strip() == 'Total' for row in tbl.rows):
        print(f"  Table at index {tbl_idx}: Total row already present — skipped.")
        return

    # Compute column sums (across years for each class)
    col_sums = [AREA_TOTALS[site][cls] for cls in CLASSES]
    total_row = tbl.add_row()
    for i, (label, val) in enumerate(zip(
        ['Total'] + [f"{v:.2f}" for v in col_sums],
        ['Total'] + [f"{v:.2f}" for v in col_sums]
    )):
        c = total_row.cells[i] if i == 0 else total_row.cells[i]
        c.text = 'Total' if i == 0 else f"{col_sums[i-1]:.2f}"
        run = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run()
        run.bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_border(c)
    print(f"  Table {tbl_idx}: Total row added {['Total'] + [f'{v:.2f}' for v in col_sums]}")

    # Find the caption paragraph for this table and add footnote after it
    body = list(doc.element.body)
    for i, child in enumerate(body):
        tag  = child.tag.split('}')[-1]
        text = ''.join(child.itertext()) if tag == 'p' else ''
        if tag == 'p' and all(k in text for k in caption_keywords):
            from docx.text.paragraph import Paragraph
            note_p = OxmlElement('w:p')
            child.addnext(note_p)
            note_para = Paragraph(note_p, doc)
            note_run  = note_para.add_run(
                "Note: minor rounding differences may occur due to pixel boundary effects.")
            note_run.italic = True
            note_run.font.size = Pt(8)
            note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f"  Footnote added after caption: {text[:60]}")
            break

add_total_row_and_footnote(2, 'keta', ('Table 3', 'Keta'))
add_total_row_and_footnote(4, 'muni', ('Table 5', 'Muni'))

# ════════════════════════════════════════════════════════════════════════════
# FIX 7 — Verify study area map is embedded
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 7] Verifying study area map embedding...")
_, cap_para = find_para('Figure 1', 'Location of Keta Lagoon Complex')
if cap_para is not None:
    # Use XML sibling lookup (avoids stale list index after earlier mutations)
    prev_elem = cap_para._p.getprevious()
    if prev_elem is not None:
        prev_xml = etree.tostring(prev_elem, encoding='unicode')
        has_drawing = 'drawing' in prev_xml
        if has_drawing:
            print("  Study area map correctly embedded as image. OK")
        else:
            print("  Image not found — re-inserting study_area_map.png...")
            smap = OUT / "study_area_map.png"
            if smap.exists():
                for child in list(prev_elem): prev_elem.remove(child)
                from docx.text.paragraph import Paragraph as _Para
                prev_para = _Para(prev_elem, doc)
                run = prev_para.add_run()
                run.add_picture(str(smap), width=Cm(14))
                prev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print("  Re-inserted.")
            else:
                print("  [WARN] study_area_map.png not found.")
    else:
        print("  [WARN] No sibling before Figure 1 caption.")
else:
    print("  [WARN] Figure 1 caption not found.")

# ════════════════════════════════════════════════════════════════════════════
# FIX 8 — Final formatting pass
# ════════════════════════════════════════════════════════════════════════════
print("[FIX 8] Final formatting pass...")

# 8a — Bold all table header rows
header_fixed = 0
for t in doc.tables:
    if not t.rows: continue
    hdr = t.rows[0]
    for cell in hdr.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if not run.bold:
                    run.bold = True
                    header_fixed += 1
            if not para.runs:
                run = para.add_run(cell.text)
                run.bold = True
                header_fixed += 1
print(f"  Table headers bolded: {header_fixed} runs updated across {len(doc.tables)} tables.")

# 8b — Italic all figure captions (paragraphs starting with "Figure" or "Figure:")
caption_fixed = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('Figure') and ':' in t[:15] or t.startswith('Figure 1.') or t.startswith('Figure:'):
        for run in p.runs:
            if not run.italic:
                run.italic = True
                caption_fixed += 1
        if not p.runs and t:
            run = p.add_run(t)
            run.italic = True
            caption_fixed += 1
print(f"  Figure captions italicised: {caption_fixed} runs updated.")

# 8c — Verify references are alphabetical
print("  Checking references order...")
refs_idx, refs_heading = find_para('5. References')
if refs_heading is None:
    refs_idx, refs_heading = find_para('References')
if refs_heading is not None:
    all_paras = doc.paragraphs
    ref_start = refs_idx + 1
    ref_texts = []
    for p in all_paras[ref_start:]:
        t = p.text.strip()
        if not t or p.style.name.startswith('Heading'): break
        ref_texts.append(t)
    sorted_texts = sorted(ref_texts, key=lambda r: r.split(',')[0].upper())
    if ref_texts == sorted_texts:
        print(f"  References already alphabetical ({len(ref_texts)} entries). OK")
    else:
        print(f"  [WARN] References not in alphabetical order — re-sorting.")
        # Remove and re-insert in order
        for p in all_paras[ref_start: ref_start + len(ref_texts)]:
            p._p.getparent().remove(p._p)
        refs_heading = doc.paragraphs[refs_idx]
        prev_p = refs_heading
        for ref in sorted_texts:
            new_p_elem = OxmlElement('w:p')
            prev_p._p.addnext(new_p_elem)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p_elem, doc)
            new_para.add_run(ref)
            new_para.paragraph_format.first_line_indent = Cm(-0.5)
            new_para.paragraph_format.left_indent       = Cm(0.5)
            prev_p = new_para
        print(f"  References re-sorted.")

# 8d — Check for duplicate consecutive paragraphs
dup_count = 0
all_texts = [p.text.strip() for p in doc.paragraphs]
seen_consecutive = set()
i = 0
paras_list = list(doc.paragraphs)
while i < len(paras_list) - 1:
    t1 = paras_list[i].text.strip()
    t2 = paras_list[i+1].text.strip()
    if t1 and t1 == t2 and len(t1) > 20:
        print(f"  [DUP] Removing duplicate paragraph: '{t1[:60]}...'")
        paras_list[i+1]._p.getparent().remove(paras_list[i+1]._p)
        paras_list = list(doc.paragraphs)
        dup_count += 1
    else:
        i += 1
print(f"  Duplicate paragraphs removed: {dup_count}")

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(str(DST_DOCX))
size_kb = DST_DOCX.stat().st_size / 1024
print(f"\n{'='*70}")
print(f"SAVED: {DST_DOCX.name}  ({size_kb:.1f} KB)")
print("=" * 70)
