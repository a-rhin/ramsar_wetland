#!/usr/bin/env python3
"""
fix_add_prop_tables.py — Insert Shannon proportion tables (9a / 10a)
immediately before landscape metrics tables (9b / 10b) in the FINAL_v2 docx.
"""
import os
os.environ['PYTHONUTF8'] = '1'

import re
import shutil
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path("D:/ramsar_wetlands")
OUT  = BASE / "outputs"
SRC  = OUT / "wetland_degradation_report_FINAL_v2.docx"
DST  = OUT / "wetland_degradation_report_FINAL_v3.docx"

shutil.copy2(str(SRC), str(DST))
doc = Document(str(DST))

YEARS   = [1991, 2001, 2015, 2025]
CLASSES = ['Vegetation', 'Water body', 'Built up/Bareland']

PI      = 'ᵢ'          # subscript i: ᵢ
KM2     = 'km²'       # km²
DASH    = '—'          # em dash —
MINUS   = '−'          # minus sign −
SIGMA   = 'Σ'          # Σ
TIMES   = '×'          # ×
DIV     = '÷'          # ÷

# ─── Compute class proportions from area CSVs ────────────────────────────────
def load_props(site):
    rows = []
    for year in YEARS:
        csv_p = BASE / f"data/{site}/area/{site}_area_{year}.csv"
        df    = pd.read_csv(csv_p)
        a_col = next(c for c in df.columns if 'Area' in c)
        n_col = next(c for c in df.columns if 'Class_name' in c)
        df2   = df[[n_col, a_col]].rename(columns={n_col: 'Class', a_col: 'Area'})
        df2['Class'] = (df2['Class'].str.strip()
                        .replace({'Built up':  'Built up/Bareland',
                                  'Built Up':  'Built up/Bareland',
                                  'Builtup':   'Built up/Bareland'}))
        total = df2['Area'].sum()
        row   = {'year': year, 'total': round(total, 2)}
        for cls in CLASSES:
            vals = df2[df2['Class'] == cls]['Area'].values
            area = float(vals[0]) if len(vals) else 0.0
            row[cls] = round(area / total, 4) if total else 0.0
        rows.append(row)
    return rows

keta_props = load_props('keta')
muni_props = load_props('muni')
print("Proportions loaded.")
for r in keta_props:
    veg = r['Vegetation']; wat = r['Water body']; blt = r['Built up/Bareland']
    print(f"  Keta {r['year']}: total={r['total']}, "
          f"veg={veg:.4f}, wat={wat:.4f}, blt={blt:.4f}, sum={veg+wat+blt:.4f}")
for r in muni_props:
    veg = r['Vegetation']; wat = r['Water body']; blt = r['Built up/Bareland']
    print(f"  Muni {r['year']}: total={r['total']}, "
          f"veg={veg:.4f}, wat={wat:.4f}, blt={blt:.4f}, sum={veg+wat+blt:.4f}")

# ─── Capture landscape metrics table element references BEFORE any mutations ──
keta_lm_tbl  = doc.tables[8]._tbl   # Table 9  — Keta landscape metrics
muni_lm_tbl  = doc.tables[9]._tbl   # Table 10 — Muni landscape metrics

# ─── Step 1: Rename Table 9 → 9b and Table 10 → 10b throughout document ─────
# Pattern: match "Table 9" or "Table 10" NOT immediately followed by a letter/digit
RENAME_RE = re.compile(r'Table (9|10)(?![a-zA-Z0-9])')

def rename_runs(para):
    for run in para.runs:
        if 'Table 9' in run.text or 'Table 10' in run.text:
            run.text = RENAME_RE.sub(lambda m: f"Table {m.group(1)}b", run.text)

renamed_count = 0
for para in doc.paragraphs:
    old = para.text
    rename_runs(para)
    if para.text != old:
        renamed_count += 1
        print(f"  Renamed para: {para.text[:80].encode('ascii','replace').decode()}")

# Also check table cells (safety pass)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                rename_runs(para)

print(f"Renaming done. {renamed_count} paragraphs updated.")

# ─── XML helpers ──────────────────────────────────────────────────────────────
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

def xml_para(text, italic=False, size_pt=None, center=False):
    """Build a standalone w:p element with a single run."""
    p = OxmlElement('w:p')
    if center:
        pPr = OxmlElement('w:pPr')
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if italic:
        r_i = OxmlElement('w:i')
        rPr.append(r_i)
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(XML_SPACE, 'preserve')
    r.append(t)
    p.append(r)
    return p

def xml_cell(text, bold=False):
    """Build a w:tc element with centred text."""
    tc  = OxmlElement('w:tc')
    p   = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc  = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b   = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = str(text)
    t.set(XML_SPACE, 'preserve')
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc

def xml_prop_table(props_data):
    """Build a w:tbl XML element for the proportion table."""
    tbl = OxmlElement('w:tbl')

    # Table properties: Table Grid style, auto width
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    # Header row
    HEADERS = ['Year', f'Total Area ({KM2})',
               f'Veg p{PI}', f'Wat p{PI}', f'Blt p{PI}']
    tr = OxmlElement('w:tr')
    for h in HEADERS:
        tr.append(xml_cell(h, bold=True))
    tbl.append(tr)

    # Data rows
    for row in props_data:
        tr = OxmlElement('w:tr')
        for val in [
            str(row['year']),
            f"{row['total']:.2f}",
            f"{row['Vegetation']:.4f}",
            f"{row['Water body']:.4f}",
            f"{row['Built up/Bareland']:.4f}",
        ]:
            tr.append(xml_cell(val))
        tbl.append(tr)

    return tbl

# ─── Step 2: Build and insert proportion tables ───────────────────────────────
FOOTNOTE = (f"p{PI} = proportion of class i = class area {DIV} total classified area. "
            f"SDI = {MINUS}{SIGMA}(p{PI} {TIMES} ln p{PI}).")

CONFIGS = [
    (
        keta_lm_tbl,
        keta_props,
        f"Table 9a: Class proportions (p{PI}) used for SDI calculation {DASH} Keta Lagoon Complex.",
    ),
    (
        muni_lm_tbl,
        muni_props,
        f"Table 10a: Class proportions (p{PI}) used for SDI calculation {DASH} Muni-Pomadze.",
    ),
]

for lm_tbl_elem, props_data, caption_text in CONFIGS:
    # Build table XML
    prop_tbl_xml = xml_prop_table(props_data)

    # Insert sequence immediately before the landscape metrics table:
    # [prop_tbl] [caption_p] [footnote_p] [empty_p] | lm_tbl
    # Add in reverse order (each addprevious inserts before the current element)

    empty_p    = xml_para('')
    lm_tbl_elem.addprevious(empty_p)           # ... | empty_p | lm_tbl

    footnote_p = xml_para(FOOTNOTE, italic=True, size_pt=8)
    empty_p.addprevious(footnote_p)            # ... | footnote_p | empty_p | lm_tbl

    caption_p  = xml_para(caption_text, italic=True, size_pt=9)
    footnote_p.addprevious(caption_p)          # ... | caption_p | footnote_p | empty_p | lm_tbl

    caption_p.addprevious(prop_tbl_xml)        # ... | prop_tbl | caption_p | footnote_p | empty_p | lm_tbl

    safe_cap = caption_text.encode('ascii', 'replace').decode()
    print(f"Inserted: {safe_cap[:60]}...")

# ─── Step 3: Verify final table count and order ───────────────────────────────
print(f"\nTotal tables in doc now: {len(doc.tables)}")
body = list(doc.element.body)
for i, tbl in enumerate(doc.tables):
    pos  = body.index(tbl._tbl)
    # find caption: look at next few siblings for a para starting with "Table"
    cap_text = ''
    for j in range(pos+1, min(pos+5, len(body))):
        txt = ''.join(body[j].itertext()).strip()
        if txt.startswith('Table'):
            cap_text = txt[:60]
            break
    safe = cap_text.encode('ascii', 'replace').decode()
    print(f"  tables[{i}] body_pos={pos}: {safe}")

# ─── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(str(DST))
size_kb = DST.stat().st_size / 1024
print(f"\nSaved: {DST.name}  ({size_kb:.1f} KB)")
